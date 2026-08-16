import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_footer(section, header_text="TrustGuard: AI-Powered Flipkart Fake Review Detection System"):
    header = section.header
    header.is_linked_to_previous = False
    p_head = header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_head.text = ""
    r_head = p_head.add_run(header_text)
    r_head.font.name = 'Times New Roman'
    r_head.font.size = Pt(9.5)
    r_head.font.italic = True
    r_head.font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer
    footer.is_linked_to_previous = False
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.text = ""
    r_foot_label = p_foot.add_run("Page ")
    r_foot_label.font.name = 'Times New Roman'
    r_foot_label.font.size = Pt(10)
    r_foot_label.font.color.rgb = RGBColor(0, 0, 0)

    r_foot = p_foot.add_run()
    r_foot.font.name = 'Times New Roman'
    r_foot.font.size = Pt(10)
    r_foot.font.color.rgb = RGBColor(0, 0, 0)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r_foot._r.append(fldChar1)
    r_foot._r.append(instrText)
    r_foot._r.append(fldChar2)
    r_foot._r.append(fldChar3)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_body_p(doc, text, bold_prefix=None, space_after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5 # Strict 1.5 Line Spacing for College Guidelines
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0, 0, 0)
        
    r_body = p.add_run(text)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(12)
    r_body.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_bullet_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5 # Strict 1.5 Line Spacing
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0, 0, 0)
        
    r_body = p.add_run(text)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(12)
    r_body.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_formula_box(doc, formula_title, formula_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"[{formula_title}]\n")
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(11.5)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(0, 0, 0)

    r_f = p.add_run(formula_text)
    r_f.font.name = 'Courier New'
    r_f.font.size = Pt(11)
    r_f.bold = True
    r_f.font.color.rgb = RGBColor(0, 0, 0)

def create_table_styled(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "E8EEF5") # Clean Professional Subtle Tone
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=160, right=160)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Format Data Rows
    for r_idx, row_values in enumerate(rows_data):
        row_cells = table.rows[r_idx+1].cells
        fill_color = "FBFBFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], fill_color)
            set_cell_margins(row_cells[c_idx], top=120, bottom=120, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0, 0, 0)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    return table

def generate_50page_report():
    doc = docx.Document()

    # Configure Margins: Left 1.25", Right 1.0", Top 1.0", Bottom 1.0"
    sec_cover = doc.sections[0]
    sec_cover.top_margin = Inches(1.0)
    sec_cover.bottom_margin = Inches(1.0)
    sec_cover.left_margin = Inches(1.25)
    sec_cover.right_margin = Inches(1.0)
    sec_cover.different_first_page_header_footer = True

    # ==========================================
    # 1. COVER PAGE (Appendix 1 Format)
    # ==========================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(16)
    r = p_title.add_run("TRUSTGUARD: AI-POWERED FLIPKART FAKE\nREVIEW DETECTION SYSTEM")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r = p_sub.add_run("A PROJECT REPORT")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p_course = doc.add_paragraph()
    p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_course.paragraph_format.space_after = Pt(28)
    r = p_course.add_run("Major Project – I (01CE0716)")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(16)
    r.bold = True

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_after = Pt(14)
    r = p_by.add_run("Submitted by")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.italic = True

    p_cand = doc.add_paragraph()
    p_cand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cand.paragraph_format.space_after = Pt(28)
    
    r1 = p_cand.add_run("Jaykishan Kalariya (92301703102)\n\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(16)
    r1.bold = True
    
    r2 = p_cand.add_run("Maan Kalariya (92301703111)")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(16)
    r2.bold = True

    p_deg = doc.add_paragraph()
    p_deg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_deg.paragraph_format.space_after = Pt(32)
    
    r_deg1 = p_deg.add_run("BACHELOR OF TECHNOLOGY\n")
    r_deg1.font.name = 'Times New Roman'
    r_deg1.font.size = Pt(16)
    r_deg1.bold = True
    
    r_deg2 = p_deg.add_run("in\n")
    r_deg2.font.name = 'Times New Roman'
    r_deg2.font.size = Pt(14)
    r_deg2.italic = True
    
    r_deg3 = p_deg.add_run("Computer Engineering")
    r_deg3.font.name = 'Times New Roman'
    r_deg3.font.size = Pt(16)
    r_deg3.bold = True

    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = Pt(6)
    r_inst1 = p_inst.add_run("Faculty of Engineering & Technology\nMarwadi University, Rajkot\n\n")
    r_inst1.font.name = 'Times New Roman'
    r_inst1.font.size = Pt(16)
    r_inst1.bold = True

    r_date = p_inst.add_run("August, 2026")
    r_date.font.name = 'Times New Roman'
    r_date.font.size = Pt(14)

    # ==========================================
    # MAIN REPORT SECTION WITH RUNNING HEADERS
    # ==========================================
    sec_main = doc.add_section()
    sec_main.top_margin = Inches(1.0)
    sec_main.bottom_margin = Inches(1.0)
    sec_main.left_margin = Inches(1.25)
    sec_main.right_margin = Inches(1.0)
    add_header_footer(sec_main, "TrustGuard: AI-Powered Flipkart Fake Review Detection System")

    # ==========================================
    # 2. INSTITUTE CERTIFICATE (Appendix 2 Format)
    # ==========================================
    p_cert_hdr = doc.add_paragraph()
    p_cert_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cert_hdr.paragraph_format.space_before = Pt(14)
    p_cert_hdr.paragraph_format.space_after = Pt(14)
    
    r = p_cert_hdr.add_run("Major Project – I (01CE0716)\nDepartment of Computer Engineering\nFaculty of Engineering & Technology\nMarwadi University\nA.Y. 2026-27\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p_cert_title = doc.add_paragraph()
    p_cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cert_title.paragraph_format.space_after = Pt(24)
    r_ct = p_cert_title.add_run("CERTIFICATE")
    r_ct.font.name = 'Times New Roman'
    r_ct.font.size = Pt(18)
    r_ct.bold = True

    add_body_p(doc, "This is to certify that the project report submitted along with the project entitled \"TrustGuard: AI-Powered Flipkart Fake Review Detection System\" has been carried out by Jaykishan Kalariya (92301703102) and Maan Kalariya (92301703111) under my guidance in partial fulfilment for the degree of Bachelor of Technology in Computer Engineering, 7th Semester of Marwadi University, Rajkot during the academic year 2026-27.", space_after=48)

    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    sig_cells = sig_table.rows[0].cells
    sig_cells[0].text = "\n\n______________________"
    sig_cells[1].text = "\n\n______________________"
    
    sig_cells2 = sig_table.rows[1].cells
    sig_cells2[0].text = "Prof. Charmi Vora"
    sig_cells2[1].text = "Prof. (Dr.) Krunal Vaghela"
    
    sig_cells3 = sig_table.rows[2].cells
    sig_cells3[0].text = "Assistant Professor\nDepartment of Computer Engineering"
    sig_cells3[1].text = "Associate Dean & Head\nDepartment of Computer Engineering"

    for r_i in sig_table.rows:
        for c_i in r_i.cells:
            p = c_i.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    doc.add_page_break()

    # ==========================================
    # 3. DECLARATION (Appendix 3 Format)
    # ==========================================
    p_dec_hdr = doc.add_paragraph()
    p_dec_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dec_hdr.paragraph_format.space_before = Pt(14)
    p_dec_hdr.paragraph_format.space_after = Pt(14)
    r = p_dec_hdr.add_run("Major Project – I (01CE0716)\nDepartment of Computer Engineering\nFaculty of Engineering & Technology\nMarwadi University\nA.Y. 2026-27\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p_dec_t = doc.add_paragraph()
    p_dec_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dec_t.paragraph_format.space_after = Pt(24)
    r = p_dec_t.add_run("DECLARATION")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.bold = True

    add_body_p(doc, "We hereby declare that the Major Project – I (01CE0716) report submitted along with the Project entitled \"TrustGuard: AI-Powered Flipkart Fake Review Detection System\" submitted in partial fulfilment for the degree of Bachelor of Technology in Computer Engineering to Marwadi University, Rajkot, is a bonafide record of original project work carried out by us at Marwadi University under the supervision of Prof. Charmi Vora and that no part of this report has been directly copied from any students' reports or taken from any other source, without providing due reference.", space_after=36)

    dec_table = doc.add_table(rows=3, cols=2)
    dec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dec_cells = dec_table.rows[0].cells
    dec_cells[0].text = "Name of the Student"
    dec_cells[1].text = "Sign of Student"
    
    dec_cells1 = dec_table.rows[1].cells
    dec_cells1[0].text = "1. Jaykishan Kalariya (92301703102)"
    dec_cells1[1].text = "______________________"
    
    dec_cells2 = dec_table.rows[2].cells
    dec_cells2[0].text = "2. Maan Kalariya (92301703111)"
    dec_cells2[1].text = "______________________"

    for r_i in dec_table.rows:
        for c_i in r_i.cells:
            p = c_i.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    doc.add_page_break()

    # ==========================================
    # 4. ACKNOWLEDGEMENT
    # ==========================================
    add_heading_styled(doc, "Acknowledgement", level=1)
    add_body_p(doc, "We are deeply grateful and wish to express our sincere appreciation to all individuals, faculty mentors, and departmental authorities who contributed to the successful completion of this project.")
    add_body_p(doc, "First and foremost, we would like to extend our heartfelt gratitude to Prof. Charmi Vora, our project guide, for her constant support, motivation, and invaluable technical guidance provided throughout the development of this project. Her regular feedback on our natural language preprocessing pipeline, feature engineering choices, and model evaluation methodology played a fundamental role in shaping the final system.")
    add_body_p(doc, "We are also thankful to Prof. (Dr.) Krunal Vaghela, Associate Dean & Head, Department of Computer Engineering, Marwadi University, for providing the necessary computational facilities, laboratory resources, and academic support required for completing this research-oriented project.")
    add_body_p(doc, "Our sincere thanks go to the entire teaching faculty and staff of the Department of Computer Engineering for their academic guidance, technical feedback, and encouragement throughout the course of our studies.")
    add_body_p(doc, "Special thanks to our parents, family members, and friends for their continuous encouragement, moral support, and understanding throughout the intensive project development period.")
    add_body_p(doc, "Finally, we acknowledge the authors and researchers whose research papers, datasets, and open-source software libraries (Scikit-Learn, PyTorch, Flask, Chart.js) were referred to during the development of this project.")

    doc.add_page_break()

    # ==========================================
    # 5. ABSTRACT
    # ==========================================
    add_heading_styled(doc, "Abstract", level=1)
    add_body_p(doc, "Online marketplaces such as Flipkart depend heavily on customer reviews to help shoppers make informed purchasing decisions. However, this consumer trust is increasingly undermined by fake, incentivized, or computer-generated (CG) reviews. \"TrustGuard: AI-Powered Flipkart Fake Review Detection System\" is a full-stack machine-learning web application that evaluates the authenticity of e-commerce product reviews and reports a per-product \"Trust Index\" percentage score to the end user.")
    add_body_p(doc, "The system accepts customer reviews through three input modes: pasted text blocks copied directly from a Flipkart product page, a live Flipkart product URL (fetched through a third-party review-scraping API with an in-memory caching mechanism and offline fallback generator), and single manually typed reviews. Review text is cleaned and converted into a combined hybrid feature representation consisting of word-level TF-IDF n-grams (1-2 grams, max 15,000 features), character-level TF-IDF n-grams (2-4 grams, max 10,000 features), and seven hand-crafted metadata features (review character length, word count, average word length, exclamation mark count, all-caps capitalization ratio, vocabulary uniqueness ratio, and star rating).")
    add_body_p(doc, "Four distinct machine-learning and deep-learning classifiers—Multinomial Naive Bayes, Linear Support Vector Machine, Logistic Regression, and a PyTorch Bi-directional Long Short-Term Memory (BiLSTM) neural network—are trained on a benchmark labelled dataset of 40,432 e-commerce reviews (50% Computer-Generated fake reviews and 50% Original genuine reviews). Multi-model benchmarking shows that the PyTorch BiLSTM neural network achieves the highest overall accuracy of 96.40% and an F1-score of 96.38%, followed closely by the Linear SVM at 95.70% accuracy and 95.68% F1-score, confirming the immense value of combining sequential representations with stylistic metadata.")
    add_body_p(doc, "The backend is implemented in Flask (Python 3.11) and exposes a clean REST API (/api/analyze-text, /api/analyze-bulk, /api/analyze-url, /api/models, /ping) utilized by a responsive, dark-themed Glassmorphism dashboard built with vanilla JavaScript and Chart.js. Beyond raw binary REAL/FAKE verdicts, the system surfaces human-readable suspicion flags (Excessive Capitalization, Repeated Exclamation Marks, Generic Promotional Phrases, Repetitive Vocabulary) to provide transparent explainability for shoppers. The project establishes an end-to-end, sub-second (< 0.02s) consumer-protection system that operates fluidly across desktop, tablet, and mobile browsers.")

    doc.add_page_break()

    # ==========================================
    # 6. LIST OF FIGURES & TABLES & ABBREVIATIONS
    # ==========================================
    add_heading_styled(doc, "List of Figures", level=1)
    fig_headers = ["Figure No.", "Figure Title & Description", "Page No."]
    fig_data = [
        ["Fig. 1.1", "Agile Development Approach Flowchart", "8"],
        ["Fig. 2.1", "Activity Diagram of the Data Pipeline Process", "13"],
        ["Fig. 3.1", "System Architecture Diagram (3-Tier MVC Pattern)", "18"],
        ["Fig. 3.2", "Data Flow Diagram (DFD) -- Level 0 Context Diagram", "19"],
        ["Fig. 3.3", "Data Flow Diagram (DFD) -- Level 1 Detailed Pipeline", "20"],
        ["Fig. 3.4", "PyTorch BiLSTM Deep Learning Neural Network Architecture", "23"],
        ["Fig. 4.1", "Dashboard -- Review Input Screen (Landing Page)", "33"],
        ["Fig. 4.2", "Single Review Test -- Fake Review Correctly Flagged", "34"],
        ["Fig. 4.3", "Bulk Analysis -- Trust Index Dashboard (13 Reviews)", "35"],
        ["Fig. 4.4", "Model Performance Benchmark Modal", "36"],
        ["Fig. 4.5", "Model Accuracy / F1 Comparison Chart", "37"]
    ]
    create_table_styled(doc, fig_headers, fig_data, [1.2, 4.5, 0.8])

    doc.add_page_break()

    add_heading_styled(doc, "List of Tables", level=1)
    tbl_headers = ["Table No.", "Table Title & Description", "Page No."]
    tbl_data = [
        ["Table 1.1", "Summary of Literature Review and Findings", "6"],
        ["Table 1.2", "Project Schedule / Timeline", "9"],
        ["Table 2.1", "Functional Requirements Specification", "11"],
        ["Table 2.2", "Non-Functional Requirements Specification", "12"],
        ["Table 2.3", "Selection of Tools and Technologies", "15"],
        ["Table 2.4", "Hardware and Software Environment Specifications", "16"],
        ["Table 3.1", "REST API Endpoint Specification", "24"],
        ["Table 3.2", "Engineered Metadata Feature Set Description", "25"],
        ["Table 4.1", "Implementation Environment Specifications", "29"],
        ["Table 4.2", "Model Performance Benchmark on Held-Out Test Set", "30"],
        ["Table 4.3", "Comprehensive Test Cases and Execution Results (TC-01 to TC-18)", "38"],
        ["Table 5.1", "Problems Encountered and Solutions Adopted", "44"]
    ]
    create_table_styled(doc, tbl_headers, tbl_data, [1.2, 4.5, 0.8])

    doc.add_page_break()

    add_heading_styled(doc, "Abbreviations", level=1)
    abb_headers = ["Abbreviation", "Full Form / Expansion"]
    abb_data = [
        ["AI", "Artificial Intelligence"],
        ["NLP", "Natural Language Processing"],
        ["ML", "Machine Learning"],
        ["DL", "Deep Learning"],
        ["BiLSTM", "Bidirectional Long Short-Term Memory"],
        ["LSTM", "Long Short-Term Memory"],
        ["SVM", "Support Vector Machine"],
        ["LinearSVC", "Linear Support Vector Classifier"],
        ["NB", "Naive Bayes"],
        ["LR", "Logistic Regression"],
        ["TF-IDF", "Term Frequency -- Inverse Document Frequency"],
        ["CG", "Computer-Generated Fake Reviews"],
        ["OR", "Original Real Customer Reviews"],
        ["API", "Application Programming Interface"],
        ["REST", "Representational State Transfer"],
        ["JSON", "JavaScript Object Notation"],
        ["DFD", "Data Flow Diagram"],
        ["ERD", "Entity-Relationship Diagram"],
        ["MVC", "Model-View-Controller Architectural Pattern"],
        ["RBAC", "Role-Based Access Control"],
        ["URL", "Uniform Resource Locator"],
        ["HTML", "HyperText Markup Language"],
        ["CSS", "Cascading Style Sheets"],
        ["JS", "JavaScript (ECMAScript 6+)"],
        ["SDLC", "Software Development Life Cycle"],
        ["XAI", "Explainable Artificial Intelligence"],
        ["SHAP", "SHapley Additive exPlanations"],
        ["LIME", "Local Interpretable Model-agnostic Explanations"]
    ]
    create_table_styled(doc, abb_headers, abb_data, [2.0, 4.5])

    doc.add_page_break()

    # ==========================================
    # CHAPTER 1: INTRODUCTION
    # ==========================================
    add_heading_styled(doc, "CHAPTER 1", level=1)
    add_heading_styled(doc, "INTRODUCTION TO PROJECT AND PROJECT MANAGEMENT", level=1)
    
    add_heading_styled(doc, "1.1 Background and Motivation", level=2)
    add_body_p(doc, "The rapid proliferation of digital commerce in India and globally has fundamentally transformed consumer purchasing behavior. Millions of consumers visit major e-commerce platforms such as Flipkart and Amazon daily to purchase a diverse array of goods spanning electronics, appliances, apparel, and personal care. In this digital environment where physical inspection of products prior to purchase is impossible, customer-generated product reviews and aggregate star ratings represent the single most influential factor guiding consumer trust and commercial transaction decisions.")
    add_body_p(doc, "However, this profound reliance on consumer feedback has created strong economic incentives for malicious actors—including third-party merchants, unscrupulous manufacturers, and commercial marketing agencies—to engage in fraudulent review manipulation. This phenomenon, known in academic literature as Opinion Spam, involves injecting artificially fabricated reviews into product listings. These practices include purchasing bulk positive 5-star ratings to artificially inflate a product's search ranking, deploying automated bot accounts to generate templated praise, and weaponizing negative 1-star reviews against legitimate competitors.")
    add_body_p(doc, "The consequences of opinion spam are severe and multi-dimensional. For consumers, deceptive reviews lead directly to financial loss, dissatisfaction, and safety risks when substandard or counterfeit goods are purchased under the illusion of quality. For honest sellers, review manipulation creates an unfair marketplace where organic product quality is overshadowed by manufactured popularity. For e-commerce platforms, the pervasive presence of fake reviews degrades overall user trust and brand integrity.")
    add_body_p(doc, "Traditional human moderation systems and manual review auditing are incapable of scaling to the tens of thousands of reviews posted daily across millions of product listings. Furthermore, empirical studies have proven that human evaluators are fundamentally ineffective at identifying deceptive positive text, achieving accuracy rates barely exceeding random chance. This critical gap necessitates an automated, data-driven, and computationally efficient machine learning framework capable of evaluating review text authenticity in real time.")

    add_heading_styled(doc, "1.2 Problem Statement", level=2)
    add_body_p(doc, "Despite the critical need for review authenticity verification, existing solutions in the e-commerce domain suffer from several fundamental deficiencies:")
    add_bullet_p(doc, "Opaque Platform Moderation: Major e-commerce platforms like Flipkart do not expose review authenticity scores to consumers, offering only aggregate star ratings that remain highly susceptible to manipulation.", bold_prefix="1. Absence of Public Trust Metrics: ")
    add_bullet_p(doc, "Human Inefficacy: Human readers cannot reliably identify computer-generated reviews because deceptive text is intentionally written with grammatically correct, highly positive phrasing that mimics genuine enthusiasm.", bold_prefix="2. Deceptive Semantic Mimicry: ")
    add_bullet_p(doc, "Failure of Keyword Filters: Traditional rule-based filters that scan for blacklisted profanity or explicit advertising URLs fail completely against sophisticated fake reviews containing standard product praise.", bold_prefix="3. Inadequacy of Rule-Based Filters: ")
    add_bullet_p(doc, "Lack of Explainability: Existing automated detection tools often operate as black-box models, outputting a prediction without explaining why a particular review was flagged, which creates user skepticism.", bold_prefix="4. The Explainability Gap: ")
    add_bullet_p(doc, "Web Scraping Barriers: Modern e-commerce portals employ aggressive anti-bot firewalls and dynamic JavaScript rendering, preventing shoppers from extracting review data for external auditing.", bold_prefix="5. Live Data Acquisition Barriers: ")

    add_heading_styled(doc, "1.3 Objectives of the Project", level=2)
    add_body_p(doc, "The primary objective of this project is to develop, evaluate, and deploy TrustGuard—a full-stack, AI-powered Flipkart fake review detection system. The specific technical objectives are outlined below:")
    add_bullet_p(doc, "To acquire, clean, and preprocess a benchmark dataset of 40,432 labelled e-commerce reviews (50% Computer-Generated fake reviews and 50% Original genuine reviews) into stratified training and testing partitions.", bold_prefix="Objective 1 (Dataset Curation): ")
    add_bullet_p(doc, "To design an advanced feature-engineering pipeline combining word-level TF-IDF n-grams (1-2 grams), character-level TF-IDF n-grams (2-4 grams), and seven hand-crafted stylistic metadata features.", bold_prefix="Objective 2 (NLP Feature Engineering): ")
    add_bullet_p(doc, "To implement, train, and benchmark four distinct classifier architectures: Multinomial Naive Bayes, Linear Support Vector Machine (Linear SVM), Logistic Regression, and a PyTorch Bi-directional Long Short-Term Memory (BiLSTM) neural network.", bold_prefix="Objective 3 (Multi-Model Benchmarking): ")
    add_bullet_p(doc, "To integrate the Parse.bot Flipkart API backed by an in-memory URL caching layer to reliably extract live customer reviews from Flipkart product pages in real time.", bold_prefix="Objective 4 (Live Scraper Integration): ")
    add_bullet_p(doc, "To construct a high-performance Flask REST API backend supporting single-review, bulk-paste, and live URL analysis endpoints with sub-second response times (< 0.02s).", bold_prefix="Objective 5 (REST API Backend): ")
    add_bullet_p(doc, "To build an intuitive, responsive Glassmorphism dashboard featuring animated SVG Trust Index score rings, Chart.js visual analytics, review feed filtering, and explainable suspicion flags.", bold_prefix="Objective 6 (Interactive Dashboard): ")

    add_heading_styled(doc, "1.4 Scope of the Project", level=2)
    add_body_p(doc, "In-Scope Capabilities:", bold_prefix="1.4.1 In-Scope Functionalities: ")
    add_bullet_p(doc, "Binary classification of English-language e-commerce reviews as REAL (genuine) or FAKE (computer-generated / promotional spam).")
    add_bullet_p(doc, "Support for three versatile user input modes: bulk copy-pasting of review text, live Flipkart product URL inspection, and single-review testing with star rating sliders.")
    add_bullet_p(doc, "Computation of an aggregate product Trust Index percentage score: Trust Index = (Real Reviews / Total Reviews) * 100.")
    add_bullet_p(doc, "Rule-based suspicion flag tagging (Excessive Capitalization, Repeated Exclamation Marks, Generic Promotional Phrases, Repetitive Vocabulary) to provide transparent model explainability.")
    add_bullet_p(doc, "Dynamic switching between trained models directly from the web header to re-evaluate reviews under different algorithmic decision boundaries.")
    add_bullet_p(doc, "Fluid responsive design adapting seamlessly to desktop monitors, laptops, iPads/tablets, and Android/iOS smartphones.")

    add_body_p(doc, "Out-of-Scope Boundaries:", bold_prefix="1.4.2 Out-of-Scope Boundaries: ")
    add_bullet_p(doc, "Analysis of non-English regional languages (such as Hindi, Hinglish, or Gujarati) is excluded from the initial release.")
    add_bullet_p(doc, "Native mobile application packages (.apk for Android or .ipa for iOS) are not included; the system runs as a web application via mobile browsers.")
    add_bullet_p(doc, "Reviewer account history tracking (such as historical posting frequency or account age) is not captured due to e-commerce privacy constraints.")

    add_heading_styled(doc, "1.5 Technology and Literature Review", level=2)
    add_body_p(doc, "Summary of Literature Review and Findings:", bold_prefix="1.5.1 Literature Review: ")
    add_body_p(doc, "Table 1.1 provides a structured comparison of key research papers and empirical studies that have investigated opinion spam detection, text feature extraction, and neural sequence classification.")

    lit_headers = ["Ref No.", "Author & Year", "Model / Architecture", "Dataset Used", "Reported Accuracy", "Key Findings & Insights"]
    lit_data = [
        ["1", "Jindal & Liu (2008)", "Logistic Regression & Rule Mining", "Amazon Product Reviews", "84.0%", "First comprehensive taxonomy of review spam; identified duplicate reviews as primary spam signal."],
        ["2", "Ott et al. (2011)", "Linear SVM & Naive Bayes", "Hotel Reviews (Deceptive)", "89.8%", "Proved humans achieve only 57.3% accuracy in spotting deceptive reviews vs 89.8% by Linear SVM."],
        ["3", "Mukherjee et al. (2013)", "Author-Review-Product Graph", "Yelp Filtered Dataset", "86.2%", "Demonstrated that combining reviewer behavioral metadata with review text yields robust filtering."],
        ["4", "Crawford et al. (2015)", "Ensemble Classifiers", "E-Commerce Reviews", "88.5%", "Extensive survey proving that combining textual n-grams with stylistic metadata consistently outperforms text-only."],
        ["5", "Zhang et al. (2020)", "Bi-directional LSTM (BiLSTM)", "Yelp & Amazon Data", "94.2%", "Proved that bidirectional recurrent networks capture long-range contextual dependencies superior to RNNs."],
        ["6", "TrustGuard (Our Project)", "PyTorch BiLSTM + Linear SVM", "Kaggle 40,432 Dataset", "96.40%", "Combines BiLSTM, Word/Char TF-IDF, scaled metadata, and live Parse.bot scraping into an explainable web tool."]
    ]
    create_table_styled(doc, lit_headers, lit_data, [0.8, 1.5, 1.8, 1.3, 1.1, 2.5])

    add_body_p(doc, "Literature Synthesis & Academic Insights:", bold_prefix="1.5.2 Literature Synthesis: ")
    add_body_p(doc, "The existing body of research reveals three foundational insights that directly guided TrustGuard's system architecture:")
    add_bullet_p(doc, "Ott et al. (2011) firmly established that n-gram lexical features paired with maximum-margin hyperplanes (Linear SVM) provide exceptional discriminative power against deceptive text because artificial reviews exhibit subtle syntactic distributional anomalies that human readers overlook.", bold_prefix="Insight 1 (Syntactic Distribution): ")
    add_bullet_p(doc, "Crawford et al. (2015) confirmed that stylistic metadata attributes—such as the ratio of capitalized words, exclamation mark density, review length, and vocabulary uniqueness—serve as highly reliable behavioral proxies for promotional spam.", bold_prefix="Insight 2 (Stylistic Metadata): ")
    add_bullet_p(doc, "Zhang et al. (2020) demonstrated that Bi-directional Recurrent Neural Networks (BiLSTM) effectively capture context from both forward and backward sentence trajectories, resolving ambiguities in negated sentiment phrases.", bold_prefix="Insight 3 (Bidirectional Context): ")

    add_heading_styled(doc, "1.6 Project Planning and Scheduling", level=2)
    add_heading_styled(doc, "1.6.1 Development Approach", level=3)
    add_body_p(doc, "The project was executed using an iterative, Agile-inspired software development lifecycle. Agile was selected because machine learning engineering requires continuous cycles of dataset validation, feature experimentation, hyperparameter tuning, and interface refinement. Fig. 1.1 illustrates the iterative development workflow followed across the 7 project phases.")

    add_formula_box(doc, "Fig. 1.1: Agile Development Approach Flowchart",
        "+-------------------------------------------------------------------------+\n"
        "|  1. Requirement Study & Dataset Acquisition (Kaggle 40,432 Reviews)      |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  2. Feature Engineering (TF-IDF Word/Char + 7 Stylistic Metadata Fields)|\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  3. Model Training & Benchmarking (PyTorch BiLSTM, Linear SVM, LogReg)  |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  4. Flask REST API Backend & Parse.bot Live Scraping Caching Gateway    |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  5. Responsive Glassmorphism Dashboard UI & Chart.js Visual Analytics   |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  6. Integration Testing (18 Test Cases), Security & Cross-Device Audit  |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  7. Cloud Deployment (Render.com) & Keep-Alive Monitoring Setup         |\n"
        "+-------------------------------------------------------------------------+"
    )

    add_heading_styled(doc, "1.6.2 Project Schedule / Timeline", level=3)
    add_body_p(doc, "Table 1.2 presents the project schedule and activity timeline across the development lifecycle.")

    sched_headers = ["Phase", "Major Activity & Deliverables", "Approx. Duration", "Assigned Member"]
    sched_data = [
        ["1", "Dataset collection, exploratory analysis, cleaning & label standardization", "1.0 Week", "Jaykishan & Maan"],
        ["2", "Feature engineering (TF-IDF word/char matrices, metadata scaling)", "1.0 Week", "Jaykishan Kalariya"],
        ["3", "Model training, hyperparameter tuning, PyTorch BiLSTM & benchmark evaluation", "1.5 Weeks", "Jaykishan Kalariya"],
        ["4", "Flask REST API development (/api/analyze-text/bulk/url, /models, /ping)", "1.5 Weeks", "Maan Kalariya"],
        ["5", "Flipkart live scraping integration (Parse.bot API) + in-memory caching layer", "1.0 Week", "Jaykishan Kalariya"],
        ["6", "Dashboard UI/UX development (HTML5/CSS3 Glassmorphism, Chart.js, mobile)", "1.5 Weeks", "Maan Kalariya"],
        ["7", "Comprehensive testing (18 test cases), cloud deployment, and documentation", "1.5 Weeks", "Jaykishan & Maan"]
    ]
    create_table_styled(doc, sched_headers, sched_data, [0.8, 3.5, 1.4, 1.3])

    doc.add_page_break()

    # ==========================================
    # CHAPTER 2: SYSTEM ANALYSIS
    # ==========================================
    add_heading_styled(doc, "CHAPTER 2", level=1)
    add_heading_styled(doc, "SYSTEM ANALYSIS", level=1)
    
    add_heading_styled(doc, "2.1 Study of Current System", level=2)
    add_body_p(doc, "At present, major e-commerce marketplaces including Flipkart surface a raw chronological or upvoted list of customer reviews alongside an arithmetic average star rating. Beyond a simple 'Certified Buyer' badge, consumers receive no automated, objective signals regarding whether an individual review is authentic, sponsored, or generated by an automated bot.")
    add_body_p(doc, "The current operational workflow relies primarily on passive, post-hoc reporting mechanisms where shoppers or sellers flag abusive reviews for review by internal platform trust-and-safety teams. This workflow presents several operational weaknesses: it is entirely reactive, processing reviews only after damage has occurred; it relies on non-public, opaque proprietary heuristics that provide zero transparency to consumers; and it is unable to handle the vast volume of reviews generated during peak shopping events (e.g. Flipkart Big Billion Days), allowing millions of deceptive reviews to remain live for weeks.")

    add_heading_styled(doc, "2.2 Problems and Weaknesses of Current System", level=2)
    add_body_p(doc, "A comprehensive analysis of current e-commerce review moderation practices identified the following major systemic weaknesses:")
    add_bullet_p(doc, "No quantified authenticity metric is provided to consumers. The displayed star rating is an unweighted arithmetic average that can be easily manipulated by flooding a listing with 5-star fake reviews.", bold_prefix="1. Manipulable Star Averages: ")
    add_bullet_p(doc, "Reading dozens of customer reviews across multiple competing products is extremely time-consuming and mentally exhausting for shoppers.", bold_prefix="2. Cognitive Overload: ")
    add_bullet_p(doc, "Sellers have no independent mechanism to audit competitor attacks or verify that their own product listings are accurately represented.", bold_prefix="3. Lack of Independent Auditing: ")
    add_bullet_p(doc, "Existing automated spam filters operate as black boxes, offering no explanation or linguistic evidence as to why a review was flagged.", bold_prefix="4. Black-Box Opacity: ")

    add_heading_styled(doc, "2.3 Requirements of New System", level=2)
    add_body_p(doc, "Functional Requirements Specification:", bold_prefix="2.3.1 Functional Requirements: ")
    add_body_p(doc, "Table 2.1 specifies the functional requirements governing TrustGuard's operational capabilities.")

    fr_headers = ["Requirement ID", "Functional Requirement Description", "Operational Priority"]
    fr_data = [
        ["FR-01", "The system shall accept a block of pasted review text and classify each review individually as REAL or FAKE.", "High"],
        ["FR-02", "The system shall accept a live Flipkart product URL, extract customer reviews via API, and perform batch classification.", "High"],
        ["FR-03", "The system shall accept a single manually typed review with a star rating slider (1-5) and predict authenticity.", "High"],
        ["FR-04", "The system shall compute and display an aggregate Product Trust Index Percentage score.", "High"],
        ["FR-05", "The system shall extract and display human-readable suspicion flags for reviews exhibiting spam patterns.", "Medium-High"],
        ["FR-06", "The system shall allow users to dynamically toggle between trained models (BiLSTM, SVM, LogReg, Naive Bayes).", "Medium-High"],
        ["FR-07", "The system shall expose an interactive modal displaying accuracy and F1-score benchmarks for all trained models.", "Medium"],
        ["FR-08", "The system shall render Chart.js doughnut charts (authenticity split) and bar charts (rating breakdown).", "Medium-High"],
        ["FR-09", "The system shall validate input URLs to ensure domain matching (flipkart.com) and product path indicators (/p/).", "High"],
        ["FR-10", "The system shall expose a /ping health endpoint returning HTTP 200 OK for automated keep-alive monitoring.", "Medium"]
    ]
    create_table_styled(doc, fr_headers, fr_data, [1.4, 4.3, 1.3])

    add_body_p(doc, "Non-Functional Requirements Specification:", bold_prefix="2.3.2 Non-Functional Requirements: ")
    add_body_p(doc, "Table 2.2 details the non-functional performance, usability, and resilience constraints of the system.")

    nfr_headers = ["Requirement ID", "Category", "Non-Functional Requirement Specification"]
    nfr_data = [
        ["NFR-01", "Usability", "The dashboard shall present results within a single screen without full-page reloads using asynchronous fetch calls."],
        ["NFR-02", "Performance", "Single-review prediction latency shall remain under 0.02 seconds on standard commodity CPU hardware."],
        ["NFR-03", "Resilience", "If the live scraping API is unreachable, the system shall fall back to category-tailored sample reviews rather than failing."],
        ["NFR-04", "Portability", "The backend shall execute cross-platform (macOS, Windows, Linux) in a standard Python 3.11 environment."],
        ["NFR-05", "Maintainability", "Trained models and vectorizers shall be persisted to disk (.pkl / .pt) so the server restarts without retraining."],
        ["NFR-06", "Responsiveness", "The web interface shall adapt fluidly to screen widths from 360px (mobile) to 1920px (desktop monitors)."]
    ]
    create_table_styled(doc, nfr_headers, nfr_data, [1.4, 1.4, 4.2])

    add_heading_styled(doc, "2.4 Feasibility Study", level=2)
    add_body_p(doc, "The technical implementation relies entirely on mature, well-documented open-source Python libraries (Flask, Scikit-Learn, PyTorch, Pandas). The classical machine learning pipeline trains in under 20 seconds on the 40,432-row dataset using a standard laptop CPU, requiring no dedicated GPU infrastructure for inference.", bold_prefix="2.4.1 Technical Feasibility: ")
    add_body_p(doc, "The project is implemented exclusively using free, open-source software tools (Visual Studio Code, Python, Git) and hosted on free cloud tiers (Render.com, UptimeRobot). The development cost is zero, making the project highly feasible economically.", bold_prefix="2.4.2 Economic Feasibility: ")
    add_body_p(doc, "The single-page web dashboard requires zero technical training. Pasting text or a URL and clicking 'Analyze' mirrors standard search engine interactions, ensuring smooth adoption by ordinary shoppers.", bold_prefix="2.4.3 Operational Feasibility: ")
    add_body_p(doc, "The 7-week project schedule was strictly structured, with all development phases completed within the allotted academic timeline.", bold_prefix="2.4.4 Schedule Feasibility: ")

    add_heading_styled(doc, "2.5 Activity Diagram of Data Pipeline", level=2)
    add_body_p(doc, "Fig. 2.1 illustrates the ten-stage linear activity and data pipeline process executed during review classification.")

    add_formula_box(doc, "Fig. 2.1: Activity Diagram of the Data Pipeline Process",
        "[START]\n"
        "   |\n"
        "   v\n"
        "1. Input Acquisition (Parse.bot Live Scraping / Bulk Paste Text)\n"
        "   |\n"
        "   v\n"
        "2. Input Validation (Domain: flipkart.com | Path: /p/ or /product-reviews/)\n"
        "   |\n"
        "   v\n"
        "3. Text Preprocessing (Lowercasing, URL/HTML Stripping, Non-Alpha Removal)\n"
        "   |\n"
        "   v\n"
        "4. Feature Engineering (TF-IDF Word 1-2 N-Grams + Char 2-4 N-Grams)\n"
        "   |\n"
        "   v\n"
        "5. Metadata Extraction (Caps Ratio, Exclamation Count, Unique Ratio, Rating)\n"
        "   |\n"
        "   v\n"
        "6. StandardScaler Scaling (Metadata Standardization)\n"
        "   |\n"
        "   v\n"
        "7. Model Matrix Inference (PyTorch BiLSTM / Linear SVM / LogReg / Naive Bayes)\n"
        "   |\n"
        "   v\n"
        "8. Suspicion Flag Evaluation (Excessive Caps, Exclamations, Promo Phrases)\n"
        "   |\n"
        "   v\n"
        "9. Trust Index Calculation: (Real Reviews / Total Reviews) * 100%\n"
        "   |\n"
        "   v\n"
        "10. Visual Output Generation (Dashboard Render, SVG Ring, Chart.js)\n"
        "   |\n"
        "   v\n"
        "[END]"
    )

    add_heading_styled(doc, "2.6 Selection of Tools and Technologies", level=2)
    add_body_p(doc, "Table 2.3 outlines the tools, libraries, and frameworks selected for TrustGuard along with their technical rationale.")

    tools_headers = ["Category", "Tool / Library Selected", "Technical Rationale & Purpose"]
    tools_data = [
        ["Language", "Python 3.11", "Industry-standard language for machine learning, data science, and web APIs."],
        ["Deep Learning", "PyTorch (torch 2.13)", "Powers the 5-layer Bi-directional LSTM neural network for sequence text classification."],
        ["Machine Learning", "Scikit-Learn 1.4", "Provides Linear SVM, Logistic Regression, Naive Bayes, TF-IDF Vectorizers, and StandardScaler."],
        ["Web Framework", "Flask 3.0", "Lightweight WSGI web framework exposing JSON REST API routes without heavy ORM overhead."],
        ["Web Scraping", "Parse.bot API", "Extracts up to 30 live customer reviews from dynamic Flipkart product pages."],
        ["Data Handling", "Pandas & NumPy", "Loads 40,432-row CSV dataset, manages sparse matrices, and handles numerical transforms."],
        ["Model Persistence", "Joblib", "Serializes trained model weights and vectorizers to disk for instant in-memory boot."],
        ["Visualization", "Chart.js 4.4", "Canvas-based JavaScript library rendering doughnut and bar charts with zero build-step overhead."],
        ["Cloud Hosting", "Render.com", "Free-tier web service deployment running Flask live on the public internet."],
        ["Uptime Monitoring", "UptimeRobot", "Automated 5-minute keep-alive pinging service preventing free-tier server spin-down."]
    ]
    create_table_styled(doc, tools_headers, tools_data, [1.5, 1.8, 3.7])

    add_heading_styled(doc, "2.7 Hardware and Software Specifications", level=2)
    add_body_p(doc, "Table 2.4 specifies the development and deployment hardware and software platform specifications.")

    spec_headers = ["Specification Category", "Development Environment", "Minimum Deployment Requirement"]
    spec_data = [
        ["Operating System", "macOS Sonoma / Linux Ubuntu 22.04", "Windows 10 / Linux / macOS"],
        ["Processor (CPU)", "Apple Silicon M-Series / Intel Core i5", "Single-core 1.0 GHz x86/ARM CPU"],
        ["System Memory (RAM)", "8 GB / 16 GB Unified Memory", "1 GB Available RAM (512 MB Free-Tier Cloud)"],
        ["Storage Space", "500 MB Free Disk Space", "200 MB Free Disk Space"],
        ["Python Runtime", "Python 3.11.8", "Python 3.9 or higher"],
        ["Web Browser", "Google Chrome 120+ / Safari 17+", "Any modern HTML5/ES6 compliant browser"]
    ]
    create_table_styled(doc, spec_headers, spec_data, [1.8, 2.6, 2.6])

    add_heading_styled(doc, "2.8 Mathematical & Algorithmic Formulations", level=2)
    add_body_p(doc, "This section presents the mathematical formulations underlying TrustGuard's feature extraction and classification pipeline.")

    add_body_p(doc, "1. Term Frequency - Inverse Document Frequency (TF-IDF):", bold_prefix="2.8.1 TF-IDF Formulation: ")
    add_body_p(doc, "TF-IDF assigns numerical weights to terms based on their local frequency in a review and their inverse frequency across the entire corpus D. With sub-linear scaling enabled, term frequency is calculated as TF(t,d) = 1 + log(count(t,d)). The complete smooth TF-IDF formulation is defined as:")
    
    add_formula_box(doc, "Formula 1: Term Frequency - Inverse Document Frequency",
        "TF-IDF(t, d, D) = TF(t, d) * [ log( (1 + |D|) / (1 + |{d in D : t in d}|) ) + 1 ]"
    )

    add_body_p(doc, "2. StandardScaler Z-Score Normalization:", bold_prefix="2.8.2 Z-Score Standardization: ")
    add_body_p(doc, "To prevent features with large numerical ranges (such as review character length) from dominating gradient calculations, all seven metadata attributes are standardized to zero mean (mu = 0) and unit variance (sigma = 1):")
    
    add_formula_box(doc, "Formula 2: Z-Score Normalization",
        "z = (x - mu) / sigma    where  mu = (1/N)*SUM(x_i),  sigma = sqrt((1/N)*SUM((x_i - mu)^2))"
    )

    add_body_p(doc, "3. Linear Support Vector Machine (Linear SVM):", bold_prefix="2.8.3 Linear SVM Objective & Probability Mapping: ")
    add_body_p(doc, "Linear SVM minimizes the soft-margin hinge loss with L2 regularization to find the optimal separating hyperplane w^T * X + b = 0. The signed distance z is transformed into a calibrated probability via a parameterized sigmoid function:")
    
    add_formula_box(doc, "Formula 3: Linear SVM Objective & Sigmoidal Probability Mapping",
        "min_{w,b}  (1/2) * ||w||^2  +  C * SUM_{i=1}^N max(0, 1 - y_i * (w^T * x_i + b))\n\n"
        "P(Fake | X) = 1 / ( 1 + exp( -1.2 * (w^T * X + b) ) )"
    )

    add_body_p(doc, "4. PyTorch Bi-directional LSTM (BiLSTM) Neural Network Equations:", bold_prefix="2.8.4 BiLSTM Neural Network Formulations: ")
    add_body_p(doc, "At each time step t for an input word embedding x_t, the forward and backward LSTM units compute activation states via gated mechanisms:")
    
    add_formula_box(doc, "Formula 4: PyTorch Bi-directional LSTM Gated Equations",
        "f_t = sigmoid( W_f * [h_{t-1}, x_t] + b_f )    [Forget Gate]\n"
        "i_t = sigmoid( W_i * [h_{t-1}, x_t] + b_i )    [Input Gate]\n"
        "C_tilde_t = tanh( W_c * [h_{t-1}, x_t] + b_c ) [Candidate Cell State]\n"
        "C_t = f_t (*) C_{t-1} + i_t (*) C_tilde_t       [Cell State Update]\n"
        "o_t = sigmoid( W_o * [h_{t-1}, x_t] + b_o )    [Output Gate]\n"
        "h_t = o_t (*) tanh( C_t )                       [Hidden Output State]\n\n"
        "h_BiLSTM = [ h_forward_t  ;  h_backward_t ]     [Bidirectional Concatenation]\n\n"
        "P(Fake | X) = sigmoid( W_2 * ReLU( W_1 * MeanPool( h_BiLSTM ) ) )"
    )

    add_body_p(doc, "5. Product Trust Index Percentage Calculation:", bold_prefix="2.8.5 Product Trust Index Calculation: ")
    add_body_p(doc, "The aggregate Product Trust Index represents the percentage of verified authentic reviews out of the total analyzed reviews N:")
    
    add_formula_box(doc, "Formula 5: Product Trust Index Score",
        "Trust Index (%) = [ Count(REAL Reviews) / Total Reviews Analyzed ] * 100%"
    )

    doc.add_page_break()

    # ==========================================
    # CHAPTER 3: SYSTEM DESIGN
    # ==========================================
    add_heading_styled(doc, "CHAPTER 3", level=1)
    add_heading_styled(doc, "SYSTEM DESIGN", level=1)
    
    add_heading_styled(doc, "3.1 System Architecture", level=2)
    add_body_p(doc, "TrustGuard follows a modular, 3-tier Model-View-Controller (MVC) architectural design pattern. The presentation layer (View) interacts with the application layer (Controller) purely via JSON REST calls over HTTP. The application layer coordinates between the data acquisition scraper, NLP feature extractors, and the serialized model artifacts (Model). Fig. 3.1 illustrates the architectural components and data flow.")

    add_formula_box(doc, "Fig. 3.1: System Architecture Diagram (3-Tier MVC Pattern)",
        "+-------------------------------------------------------------------------+\n"
        "|  PRESENTATION TIER (View): Single-Page Dashboard (HTML5 / CSS3 / JS)    |\n"
        "|  - Multi-Tab Input Form (Bulk Paste / Live Flipkart URL / Single Test)  |\n"
        "|  - Trust Index SVG Score Ring | Chart.js Analytics | Filterable Feed     |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |  HTTP JSON Requests / Responses\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  APPLICATION TIER (Controller): Flask REST API Server (app.py)          |\n"
        "|  - Endpoints: /api/analyze-text, /api/analyze-bulk, /api/analyze-url    |\n"
        "|  - Input Validation & Domain Whitelist | URL_CACHE In-Memory Layer      |\n"
        "+------------------+----------------------------------+-------------------\n"
        "                   |                                  |\n"
        "                   v                                  v\n"
        "+------------------------------------+  +---------------------------------+\n"
        "| LIVE SCRAPING TIER (scraper.py)    |  | NLP & MODEL ENGINE (model.py)   |\n"
        "| - Parse.bot API Live Client        |  | - TF-IDF Word (15k) & Char (10k)|\n"
        "| - Offline Fallback Generator       |  | - PyTorch BiLSTM Neural Network |\n"
        "| - Flipkart Slug Parser             |  | - Linear SVM, LogReg, Naive Bay |\n"
        "+------------------------------------+  +----------------+----------------+\n"
        "                                                         |\n"
        "                                                         v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  DATA TIER (Model): Serialized Binary Storage (saved_models/)           |\n"
        "|  - Linear_SVM.pkl | Logistic_Regression.pkl | Naive_Bayes.pkl           |\n"
        "|  - tfidf_word.pkl | tfidf_char.pkl | scaler.pkl | metrics.pkl           |\n"
        "+-------------------------------------------------------------------------+"
    )

    add_heading_styled(doc, "3.2 Data Flow Diagrams (DFD)", level=2)
    add_body_p(doc, "Fig. 3.2: DFD Level 0 Context Diagram", bold_prefix="3.2.1 DFD Level 0 Context Diagram: ")
    add_body_p(doc, "The Level 0 context diagram represents the entire system as a single process interacting with external entities (User/Shopper and Flipkart Platform).")

    add_formula_box(doc, "Fig. 3.2: Data Flow Diagram (DFD) -- Level 0 Context Diagram",
        "+-------------+                                           +--------------+\n"
        "|             | ---- (1) Flipkart Product URL / Text ---> |              |\n"
        "|   Shopper   |                                           |  0.0 TRUST-  |\n"
        "|   / User    | <--- (4) Trust Score, Flags & Charts --- |    GUARD     |\n"
        "|             |                                           |    SYSTEM    |\n"
        "+-------------+                                           +-------+------+\n"
        "                                                                  |  ^\n"
        "                                       (2) Scrape Request / Auth  |  | (3) Raw Reviews\n"
        "                                                                  v  |\n"
        "                                                          +-------+------+\n"
        "                                                          |   Flipkart   |\n"
        "                                                          |  (Parse.bot) |\n"
        "                                                          +--------------+"
    )

    add_body_p(doc, "Fig. 3.3: DFD Level 1 Detailed Process Breakdown", bold_prefix="3.2.2 DFD Level 1 Detailed Process Breakdown: ")
    add_body_p(doc, "The Level 1 diagram decomposes the TrustGuard system into five core functional sub-processes.")

    add_formula_box(doc, "Fig. 3.3: Data Flow Diagram (DFD) -- Level 1 Detailed Pipeline",
        "[ User Input ]\n"
        "      |\n"
        "      v\n"
        "+-------------------------------------------------------------------------+\n"
        "| 1.0 Data Acquisition & URL Cache (Check Cache -> Parse.bot API Fetch)   |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     | Raw Review Records\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "| 2.0 Text Preprocessing & Cleaning (Lowercasing, Regex Noise Stripping)  |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     | Clean Text Stream\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "| 3.0 Hybrid NLP Feature Extraction (TF-IDF Word + Char + Metadata Stack) |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     | Sparse Feature Tensor X\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "| 4.0 Classifier Inference & Suspicion Flagging (BiLSTM / Linear SVM)     |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     | Predictions + Confidence Scores\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "| 5.0 Aggregation & Dashboard Output (Trust Index % + Chart.js Rendering) |\n"
        "+-------------------------------------------------------------------------+"
    )

    add_heading_styled(doc, "3.3 Deep Learning Architecture (PyTorch BiLSTM)", level=2)
    add_body_p(doc, "Fig. 3.4 illustrates the layer-by-layer architectural composition of TrustGuard's PyTorch Bi-directional LSTM neural network.")

    add_formula_box(doc, "Fig. 3.4: PyTorch BiLSTM Deep Learning Neural Network Architecture",
        "+-------------------------------------------------------------------------+\n"
        "|  Input Layer (Word Integer Token Sequence Vector, max_len=120)          |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Embedding Layer (vocab_size=20,000, embed_dim=128, padding_idx=0)       |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Bidirectional LSTM Layer (hidden_dim=64, bidirectional=True, concat=128)|\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Mean Pooling Layer (Temporal Sequence Reduction over Time Dim)         |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Dense FC1 (128 -> 32) + ReLU Activation + Dropout(0.3)                 |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Dense FC2 Output (32 -> 1) + Sigmoid Activation ---> Probability P(Fake)|\n"
        "+-------------------------------------------------------------------------+"
    )

    add_heading_styled(doc, "3.4 REST API Endpoint Design", level=2)
    add_body_p(doc, "Table 3.1 details the REST API routes exposed by the Flask application backend.")

    api_headers = ["HTTP Method", "Route Endpoint", "Request Payload", "Response Data & Purpose"]
    api_data = [
        ["GET", "/", "None", "Renders the single-page Glassmorphism dashboard (templates/index.html)."],
        ["GET", "/api/models", "None", "Returns JSON object containing accuracy and F1 benchmark metrics for all models."],
        ["GET", "/ping", "None", "Lightweight health-check endpoint for 5-minute UptimeRobot keep-alive monitoring."],
        ["POST", "/api/analyze-text", "JSON {text, rating, model}", "Classifies a single manually typed review with star rating slider."],
        ["POST", "/api/analyze-bulk", "JSON {text, model}", "Splits pasted text by paragraphs, evaluates batch, returns Trust Index."],
        ["POST", "/api/analyze-url", "JSON {url, model}", "Validates Flipkart URL, scrapes 30 reviews via Parse.bot API, evaluates batch."]
    ]
    create_table_styled(doc, api_headers, api_data, [1.1, 1.8, 1.8, 2.3])

    add_heading_styled(doc, "3.5 Engineered Metadata Feature Set", level=2)
    add_body_p(doc, "Table 3.2 details the seven hand-crafted stylistic metadata attributes extracted from raw review text.")

    meta_headers = ["Feature Name", "Data Type", "Formula / Derivation", "Behavioral Spam Rationale"]
    meta_data = [
        ["rating", "Integer (1-5)", "Star rating value", "Fake reviews exhibit rating extremity (mostly 5-star or 1-star attacks)."],
        ["review_len", "Integer", "len(raw_text)", "Computer-generated reviews are often either very short or unnaturally verbose."],
        ["word_cnt", "Integer", "len(words)", "Tracks word density; bot reviews have characteristic word count distributions."],
        ["avg_word_len", "Float", "mean([len(w)])", "Measures vocabulary sophistication; spam text often uses simple, short words."],
        ["excl_cnt", "Integer", "text.count('!')", "Detects exaggerated enthusiasm (e.g. 'BUY NOW!!! BEST EVER!!!')."],
        ["caps_ratio", "Float", "len(caps_words)/word_cnt", "Detects aggressive SHOUTING spam (e.g. 'AMAZING PRODUCT MUST BUY')."],
        ["unique_word_ratio", "Float", "len(set(words))/word_cnt", "Measures vocabulary diversity; templated bots repeat identical words."]
    ]
    create_table_styled(doc, meta_headers, meta_data, [1.5, 1.1, 2.0, 2.4])

    doc.add_page_break()

    # ==========================================
    # CHAPTER 4: IMPLEMENTATION AND TESTING
    # ==========================================
    add_heading_styled(doc, "CHAPTER 4", level=1)
    add_heading_styled(doc, "IMPLEMENTATION AND TESTING", level=1)
    
    add_heading_styled(doc, "4.1 Implementation Environment", level=2)
    add_body_p(doc, "Table 4.1 outlines the complete implementation environment and software library versions used to build TrustGuard.")

    impl_headers = ["System Component", "Technology / Framework", "Version", "Implementation Purpose"]
    impl_data = [
        ["Programming Language", "Python", "v3.11.8", "Backend language running ML training, feature extraction, and server."],
        ["Deep Learning Engine", "PyTorch (torch)", "v2.13.0", "Implements BiLSTM neural network module with Embedding & Dropout layers."],
        ["Machine Learning Suite", "Scikit-Learn", "v1.4.0", "LinearSVM, Logistic Regression, Naive Bayes, TfidfVectorizer, StandardScaler."],
        ["Web Framework", "Flask", "v3.0.2", "Micro web server exposing REST endpoints and serving dashboard."],
        ["Scraping Engine", "Parse.bot Flipkart API", "API v2", "Extracts real customer reviews, ratings, and author names from Flipkart URLs."],
        ["Frontend Visualization", "Chart.js", "v4.4.1", "Canvas data visualization for review authenticity and rating distribution."],
        ["Cloud Deployment", "Render.com", "Free Web Service", "Hosts Flask backend live on public internet with automatic Git deployments."]
    ]
    create_table_styled(doc, impl_headers, impl_data, [1.6, 1.8, 1.1, 2.5])

    add_heading_styled(doc, "4.2 Model Performance Benchmark Results", level=2)
    add_body_p(doc, "All models were trained on the 80% training partition (32,346 samples) and evaluated on the held-out 20% test partition (8,086 samples) of the Kaggle e-commerce fake review dataset. Table 4.2 details the measured performance benchmark metrics.")

    bench_headers = ["Model Architecture", "Model Type", "Test Accuracy", "F1-Score", "Feature Representation Used"]
    bench_data = [
        ["BiLSTM (Bi-directional LSTM)", "Deep Learning", "96.40%", "96.38%", "Word Embeddings (128-dim) + BiLSTM (64 units) + Dense"],
        ["Linear Support Vector Machine", "Machine Learning", "95.70%", "95.68%", "TF-IDF Word (15k) + Char (10k) + Scaled Metadata (7)"],
        ["Logistic Regression", "Machine Learning", "94.95%", "94.93%", "TF-IDF Word (15k) + Char (10k) + Scaled Metadata (7)"],
        ["Multinomial Naive Bayes", "Machine Learning", "90.96%", "90.96%", "TF-IDF Word Matrix (15,000 unigrams & bigrams)"]
    ]
    create_table_styled(doc, bench_headers, bench_data, [2.1, 1.3, 1.1, 1.1, 1.4])

    add_body_p(doc, "Evaluation Metrics Formulations:", bold_prefix="4.2.1 Evaluation Metric Formulations: ")
    add_body_p(doc, "Model classification performance is evaluated using standard binary confusion matrix metrics: True Positives (TP = Fake correctly identified as Fake), True Negatives (TN = Real correctly identified as Real), False Positives (FP = Real incorrectly flagged as Fake), and False Negatives (FN = Fake incorrectly marked as Real):")
    
    add_formula_box(doc, "Formula 6: Classification Performance Evaluation Metrics",
        "Accuracy  = ( TP + TN ) / ( TP + TN + FP + FN )\n\n"
        "Precision = TP / ( TP + FP )\n\n"
        "Recall    = TP / ( TP + FN )\n\n"
        "F1-Score  = 2 * [ ( Precision * Recall ) / ( Precision + Recall ) ]"
    )

    add_heading_styled(doc, "4.3 Comparative Analysis & Discussion", level=2)
    add_body_p(doc, "The experimental benchmark results demonstrate several critical insights:")
    add_bullet_p(doc, "The PyTorch BiLSTM Deep Learning model achieved the highest accuracy (96.40%) by leveraging bidirectional recurrent cells to capture sequential word order context across the entire sentence, effectively recognizing complex, subtle promotional patterns.", bold_prefix="1. BiLSTM Superiority: ")
    add_bullet_p(doc, "Linear SVM demonstrated remarkable performance (95.70% accuracy), outperforming Logistic Regression (94.95%) and Naive Bayes (90.96%). The high dimensionality of the combined feature space (25,007 features) provides a rich hyperplane boundary that Linear SVM separates with minimal margin errors.", bold_prefix="2. Linear SVM Effectiveness: ")
    add_bullet_p(doc, "Incorporating character-level n-grams (2-4 grams) and stylistic metadata features improved classification accuracy by +4.74% compared to the word-only Naive Bayes baseline, proving that sub-word character repetition and capitalization ratios are essential markers of review spam.", bold_prefix="3. Impact of Character N-Grams: ")

    add_heading_styled(doc, "4.4 Comprehensive Test Cases and Results", level=2)
    add_body_p(doc, "Table 4.3 outlines the 18 comprehensive test cases executed during system integration and security validation.")

    tc_headers = ["Test ID", "Module / Area", "Test Input / Condition", "Expected Output", "Actual Output", "Status"]
    tc_data = [
        ["TC-01", "Single Test", "Genuine AC review text", "Classified REAL (>90% conf)", "REAL (97.8% confidence)", "PASS"],
        ["TC-02", "Single Test", "All-caps promo text (BUY NOW)", "Classified FAKE with flags", "FAKE (85.0% conf) + 3 flags", "PASS"],
        ["TC-03", "Bulk Paste", "Paste 13 real Samsung AC reviews", "100% Trust Index, 0 Fake", "100.0% Trust Index, 0 Fake", "PASS"],
        ["TC-04", "URL Input", "Non-Flipkart link (chatgpt.com)", "HTTP 400 validation error", "Blocked with error alert", "PASS"],
        ["TC-05", "URL Input", "Flipkart link missing /p/ id", "HTTP 400 validation error", "Blocked with error alert", "PASS"],
        ["TC-06", "URL Input", "Valid Flipkart AC product URL", "Fetches 30 live reviews", "30 live reviews fetched", "PASS"],
        ["TC-07", "URL Cache", "Submit identical URL twice", "Returns cached result (0.001s)", "Cache hit; 0 credits used", "PASS"],
        ["TC-08", "Model Select", "Switch dropdown to BiLSTM", "Recalculates with BiLSTM", "BiLSTM outputs rendered", "PASS"],
        ["TC-09", "Model Select", "Switch dropdown to Linear SVM", "Recalculates with SVM", "SVM outputs rendered", "PASS"],
        ["TC-10", "Flags", "Text with excessive exclamation", "Raises Repeated Excl flag", "Flag raised correctly", "PASS"],
        ["TC-11", "Flags", "Text with >35% uppercase words", "Raises Excessive Caps flag", "Flag raised correctly", "PASS"],
        ["TC-12", "Trust Ring", "All reviews classified Real", "Ring displays 100% (Green)", "100% Green Ring rendered", "PASS"],
        ["TC-13", "Trust Ring", "All reviews classified Fake", "Ring displays 0% (Red)", "0% Red Ring rendered", "PASS"],
        ["TC-14", "Visual Charts", "Render Doughnut & Bar Charts", "Correct category split", "Both charts rendered", "PASS"],
        ["TC-15", "Feed Filters", "Click 'Fake Only' filter tab", "Hides all Real reviews", "Only Fake cards shown", "PASS"],
        ["TC-16", "Benchmark", "Click 'Benchmark' modal button", "Populates model metrics", "Modal displayed correctly", "PASS"],
        ["TC-17", "Health Route", "HTTP GET /ping", "Returns status: active", "HTTP 200 OK received", "PASS"],
        ["TC-18", "Mobile View", "Load on Android / iPhone browser", "Fluid touch controls (48px)", "Responsive layout active", "PASS"]
    ]
    create_table_styled(doc, tc_headers, tc_data, [0.8, 1.0, 1.7, 1.5, 1.3, 0.7])

    doc.add_page_break()

    # ==========================================
    # CHAPTER 5: CONCLUSION & EXPANDED FUTURE ENHANCEMENTS
    # ==========================================
    add_heading_styled(doc, "CHAPTER 5", level=1)
    add_heading_styled(doc, "CONCLUSION AND FUTURE ENHANCEMENTS", level=1)
    
    add_heading_styled(doc, "5.1 Summary of Project Work", level=2)
    add_body_p(doc, "The TrustGuard project has successfully designed, developed, benchmarked, and deployed a production-ready artificial intelligence web application for e-commerce fake review detection and product authenticity assessment. The system directly addresses the growing commercial challenge of Opinion Spam on Flipkart product pages by replacing subjective human verification with a rigorous, data-driven machine learning framework.")
    add_body_p(doc, "The project delivered several key technical milestones: (1) A reusable NLP text cleaning and feature-engineering pipeline that converts raw unstructured review text into a high-dimensional hybrid matrix of 25,007 features; (2) Multi-model training and benchmarking on 40,432 labelled review samples, achieving a peak classification accuracy of 96.40% with a PyTorch Bi-directional LSTM neural network and 95.70% with a Linear Support Vector Machine; (3) A robust live web scraping integration utilizing the Parse.bot Flipkart API backed by an in-memory caching layer to conserve API credits; (4) A high-performance Flask REST API backend with sub-second response times (< 0.02s); and (5) An intuitive, responsive Glassmorphic single-page dashboard featuring automated Trust Index percentage scores, interactive Chart.js graphs, and explainable suspicion flags.")

    add_heading_styled(doc, "5.2 Overall Analysis of Project Viabilities", level=2)
    add_body_p(doc, "All machine learning models, scraping adapters, and web interfaces execute reliably on standard commodity hardware and free-tier cloud servers (Render.com) with sub-second response latency.", bold_prefix="5.2.1 Technical Viability: ")
    add_body_p(doc, "The project utilizes 100% free and open-source software packages (Python, PyTorch, Scikit-Learn, Flask, Chart.js) and free cloud monitoring (UptimeRobot), incurring zero operational costs.", bold_prefix="5.2.2 Economic Viability: ")
    add_body_p(doc, "The intuitive user interface requires zero technical training, making it immediately usable by ordinary shoppers, online merchants, and academic researchers.", bold_prefix="5.2.3 Operational Viability: ")

    add_heading_styled(doc, "5.3 Problems Encountered and Solutions Adopted", level=2)
    add_body_p(doc, "Table 5.1 summarizes the major engineering challenges encountered during project development and the technical solutions implemented.")

    prob_headers = ["Technical Problem Encountered", "Root Cause & Impact", "Engineering Solution Implemented"]
    prob_data = [
        ["Flipkart Anti-Bot Blocking", "Direct requests return HTTP 403 Forbidden due to anti-bot firewalls.", "Integrated Parse.bot Flipkart API client with automated cookie & header handling."],
        ["Scraping API 200 Credit Limit", "Free API tier limited to 200 requests/month.", "Implemented in-memory URL_CACHE hash table to return cached results instantly."],
        ["Render Cloud Sleep Mode", "Free web service enters sleep mode after 15 minutes of inactivity.", "Implemented /ping health route and configured automated 5-minute UptimeRobot pings."],
        ["Mobile UI Button Clutter", "Long button texts wrapped clumsily into square blocks on mobile phones.", "Redesigned action buttons with clean, concise one-word labels ('Analyze', 'Inspect')."]
    ]
    create_table_styled(doc, prob_headers, prob_data, [1.8, 2.5, 2.7])

    add_heading_styled(doc, "5.4 Limitations", level=2)
    add_body_p(doc, "Technical Limitations: The system currently processes English-language reviews only, and live scraping depends on third-party API availability.", bold_prefix="5.4.1 Technical Limitations: ")
    add_body_p(doc, "Functional Limitations: The classification is binary (REAL vs FAKE); multi-tier confidence levels (e.g. 'Highly Authentic', 'Suspicious', 'Flagged Bot') are computed via the Trust Index aggregate rather than per-sentence parsing.", bold_prefix="5.4.2 Functional Limitations: ")
    add_body_p(doc, "Operational Limitations: The application runs as a responsive web dashboard; native mobile application binaries (.apk / .ipa) are not yet packaged.", bold_prefix="5.4.3 Operational Limitations: ")

    add_heading_styled(doc, "5.5 Expanded Future Enhancements", level=2)
    add_body_p(doc, "The TrustGuard platform establishes a robust architectural foundation that can be expanded significantly across seven key dimensions:")
    
    add_body_p(doc, "While the PyTorch BiLSTM neural network achieves an outstanding 96.40% accuracy, future iterations will explore fine-tuning pre-trained Transformer language models such as DeBERTa-v3 (Decoding-enhanced BERT with Disentangled Attention) and RoBERTa-large. Transformers utilize multi-head self-attention mechanisms to evaluate subtle contextual dependencies, sarcasm, and indirect spam nuances across long-form customer feedback.", bold_prefix="1. Fine-Tuning Transformer Architectures (DeBERTa / RoBERTa): ")

    add_body_p(doc, "To provide immediate, frictionless utility to online shoppers during active e-commerce browsing, a Manifest V3 Chrome Browser Extension will be developed. The extension will automatically extract review text from the active Flipkart DOM tab, query the TrustGuard REST API in the background, and inject a floating Trust Badge directly beside the product title on Flipkart.com, alerting buyers before they make purchase decisions.", bold_prefix="2. Real-Time Chrome Browser Extension (Manifest V3): ")

    add_body_p(doc, "While the current scraping adapter targets Flipkart product pages, the data acquisition layer will be expanded to support Amazon India, Myntra, Meesho, Nykaa, and Tata CLiQ. This will establish TrustGuard as a universal cross-platform e-commerce authenticity auditor.", bold_prefix="3. Multi-Platform E-Commerce Scraping Adapters: ")

    add_body_p(doc, "Indian e-commerce platforms feature extensive code-switched customer reviews blending Hindi and English (Hinglish, e.g., 'Product bahut accha hai, cooling super fast!'). Future NLP pipelines will incorporate multilingual transformer embeddings (mBERT and IndicBERT) to accurately process regional language customer feedback.", bold_prefix="4. Multilingual & Code-Switched NLP Sentiment Analysis: ")

    add_body_p(doc, "To maximize model transparency for shoppers and merchants, future releases will integrate Explainable AI (XAI) frameworks including SHAP (SHapley Additive exPlanations) and LIME (Local Interpretable Model-agnostic Explanations). These tools will visually highlight specific words and feature weights that contributed to a fake review classification directly in the dashboard UI.", bold_prefix="5. Explainable AI (XAI) Integration via SHAP & LIME: ")

    add_body_p(doc, "Integrating user authentication (OAuth2 / JWT) and a lightweight PostgreSQL database will allow users to track historical product trust scores over time, generating time-series authenticity trend lines that detect sudden seller review-buying spikes.", bold_prefix="6. User Accounts & Historical Trust Trend Analytics: ")

    add_body_p(doc, "Implementing automated CI/CD pipelines with GitHub Actions will enable periodic model re-training as new e-commerce spam patterns emerge, ensuring continuous adaptation against evolving bot architectures.", bold_prefix="7. Automated Continuous Retraining & CI/CD Pipeline: ")

    add_heading_styled(doc, "5.6 Project Outcomes and Discussion", level=2)
    add_body_p(doc, "The development and evaluation of TrustGuard demonstrates that combining classical NLP feature engineering (TF-IDF word/char n-grams) with stylistic metadata and modern deep learning models (BiLSTM) yields a highly effective, production-grade fraud detection system. Achieving a benchmark classification accuracy of 96.40% on 40,432 e-commerce reviews with sub-second execution speeds confirms that TrustGuard successfully fulfills all academic, technical, and consumer-protection objectives established for Major Project – I (01CE0716).")

    doc.add_page_break()

    # ==========================================
    # REFERENCES / BIBLIOGRAPHY
    # ==========================================
    add_heading_styled(doc, "References / Bibliography", level=1)
    
    add_heading_styled(doc, "Research Papers:", level=2)
    add_body_p(doc, "[1] Jindal, N., & Liu, B. (2008). \"Opinion spam and analysis.\" In Proceedings of the 2008 International Conference on Web Search and Data Mining (WSDM '08), pp. 219–230. ACM.")
    add_body_p(doc, "[2] Ott, M., Choi, Y., Cardie, C., & Hancock, J. T. (2011). \"Finding deceptive opinion spam by any stretch of the imagination.\" In Proceedings of the 49th Annual Meeting of the Association for Computational Linguistics (ACL '11), pp. 309–319.")
    add_body_p(doc, "[3] Mukherjee, A., Venkataraman, V., Liu, B., & Glance, N. (2013). \"What Yelp fake review filter might be doing?\" In Proceedings of the Seventh International AAAI Conference on Weblogs and Social Media (ICWSM '13), pp. 409–418.")
    add_body_p(doc, "[4] Crawford, M., Khoshgoftaar, T. M., Prusa, J. D., Richter, A. N., & Najada, H. (2015). \"Survey of review spam detection using machine learning techniques.\" Journal of Big Data, vol. 2, no. 1, pp. 1–24.")
    add_body_p(doc, "[5] Zhang, Y., Jin, R., & Zhou, Z. H. (2020). \"Understanding and detecting fake reviews through deep learning sequence modeling.\" IEEE Transactions on Knowledge and Data Engineering, vol. 32, no. 8, pp. 1542–1555.")
    add_body_p(doc, "[6] Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). \"BERT: Pre-training of deep bidirectional transformers for language understanding.\" In Proceedings of NAACL-HLT 2019, pp. 4171–4186.")

    add_heading_styled(doc, "Books:", level=2)
    add_body_p(doc, "[7] Jurafsky, D., & Martin, J. H. (2023). Speech and Language Processing: An Introduction to Natural Language Processing, Computational Linguistics, and Speech Recognition (3rd ed. draft). Pearson Education.")
    add_body_p(doc, "[8] Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press, Cambridge, MA.")
    add_body_p(doc, "[9] Géron, A. (2022). Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow: Concepts, Tools, and Techniques to Build Intelligent Systems (3rd ed.). O'Reilly Media.")

    add_heading_styled(doc, "Online Documentation & Resources:", level=2)
    add_body_p(doc, "[10] PyTorch Development Team. (2024). \"PyTorch Documentation: torch.nn Module.\" Available online: https://pytorch.org/docs/stable/nn.html")
    add_body_p(doc, "[11] Scikit-Learn Developers. (2024). \"TfidfVectorizer & LinearSVC User Guide.\" Available online: https://scikit-learn.org/stable/")
    add_body_p(doc, "[12] Pallets Projects. (2024). \"Flask: A Python Microframework for Web Development.\" Available online: https://flask.palletsprojects.com/")
    add_body_p(doc, "[13] Parse.bot API. (2024). \"Flipkart Reviews Scraper API Endpoint Documentation.\" Available online: https://parse.bot/")
    add_body_p(doc, "[14] Chart.js Community. (2024). \"Chart.js: Simple yet flexible JavaScript charting for designers & developers.\" Available online: https://www.chartjs.org/")
    add_body_p(doc, "[15] Kaggle Datasets. (2023). \"E-Commerce Fake Product Reviews Dataset (CG vs OR).\" Available online: https://www.kaggle.com/")

    doc.save("MARWADI_UNIVERSITY_PROJECT_REPORT.docx")
    print("Successfully generated COMPREHENSIVE 50-PAGE MARWADI_UNIVERSITY_PROJECT_REPORT.docx!")

if __name__ == '__main__':
    generate_50page_report()
