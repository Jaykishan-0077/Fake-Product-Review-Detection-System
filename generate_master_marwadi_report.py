import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_footer(section, header_text="Department of Computer Engineering | Marwadi University"):
    header = section.header
    header.is_linked_to_previous = False
    p_head = header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_head.text = ""
    r_head = p_head.add_run(header_text)
    r_head.font.name = 'Times New Roman'
    r_head.font.size = Pt(9)
    r_head.font.italic = True
    r_head.font.color.rgb = RGBColor(120, 120, 120)

    footer = section.footer
    footer.is_linked_to_previous = False
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.text = ""
    r_foot = p_foot.add_run()
    r_foot.font.name = 'Times New Roman'
    r_foot.font.size = Pt(10)
    r_foot.font.color.rgb = RGBColor(0, 0, 0)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r_foot._r.append(fldChar1)
    r_foot._r.append(instrText)
    r_foot._r.append(fldChar2)
    r_foot._r.append(fldChar3)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_body_p(doc, text, bold_prefix=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0, 0, 0)
        
    r_body = p.add_run(text)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(12)
    r_body.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_bullet_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0, 0, 0)
        
    r_body = p.add_run(text)
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(12)
    r_body.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_formula_box(doc, formula_title, formula_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    r_title = p.add_run(f"[{formula_title}]\n")
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(11)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)

    r_f = p.add_run(formula_text)
    r_f.font.name = 'Courier New'
    r_f.font.size = Pt(11)
    r_f.bold = True
    r_f.font.color.rgb = RGBColor(150, 0, 0)

def create_table_styled(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Format Data Rows
    for r_idx, row_values in enumerate(rows_data):
        row_cells = table.rows[r_idx+1].cells
        fill_color = "F2F4F8" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], fill_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def generate_master_report():
    doc = docx.Document()

    sec_cover = doc.sections[0]
    sec_cover.top_margin = Inches(1.0)
    sec_cover.bottom_margin = Inches(1.0)
    sec_cover.left_margin = Inches(1.25)
    sec_cover.right_margin = Inches(1.0)
    sec_cover.different_first_page_header_footer = True

    # ==========================================
    # 1. COVER PAGE (Appendix 1 Format)
    # ==========================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(30)
    p_title.paragraph_format.space_after = Pt(16)
    r = p_title.add_run("TrustGuard: AI-Powered Flipkart Fake Review Detection System using Machine Learning & Deep Learning")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    r = p_sub.add_run("A PROJECT REPORT")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p_course = doc.add_paragraph()
    p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_course.paragraph_format.space_after = Pt(24)
    r = p_course.add_run("Major Project – I (01CE0716)")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(16)
    r.bold = True

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_after = Pt(12)
    r = p_by.add_run("Submitted by")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.italic = True

    p_cand = doc.add_paragraph()
    p_cand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cand.paragraph_format.space_after = Pt(24)
    
    r1 = p_cand.add_run("Jaykishan Kalariya\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(16)
    r1.bold = True
    
    r1_e = p_cand.add_run("(92301703102)\n\n")
    r1_e.font.name = 'Times New Roman'
    r1_e.font.size = Pt(14)
    
    r2 = p_cand.add_run("Maan Kalariya\n")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(16)
    r2.bold = True
    
    r2_e = p_cand.add_run("(92301703111)")
    r2_e.font.name = 'Times New Roman'
    r2_e.font.size = Pt(14)

    p_deg = doc.add_paragraph()
    p_deg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_deg.paragraph_format.space_after = Pt(28)
    
    r_deg1 = p_deg.add_run("BACHELOR OF TECHNOLOGY\n")
    r_deg1.font.name = 'Times New Roman'
    r_deg1.font.size = Pt(16)
    r_deg1.bold = True
    
    r_deg2 = p_deg.add_run("in\n")
    r_deg2.font.name = 'Times New Roman'
    r_deg2.font.size = Pt(14)
    r_deg2.italic = True
    
    r_deg3 = p_deg.add_run("Computer Engineering")
    r_deg3.font.name = 'Times New Roman'
    r_deg3.font.size = Pt(16)
    r_deg3.bold = True

    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = Pt(6)
    r_inst1 = p_inst.add_run("Faculty of Engineering & Technology\nMarwadi University, Rajkot\n")
    r_inst1.font.name = 'Times New Roman'
    r_inst1.font.size = Pt(16)
    r_inst1.bold = True

    r_date = p_inst.add_run("August, 2026")
    r_date.font.name = 'Times New Roman'
    r_date.font.size = Pt(14)

    # Main Section with Running Header & Footer
    sec_main = doc.add_section()
    sec_main.top_margin = Inches(1.0)
    sec_main.bottom_margin = Inches(1.0)
    sec_main.left_margin = Inches(1.25)
    sec_main.right_margin = Inches(1.0)
    add_header_footer(sec_main, "Department of Computer Engineering | Marwadi University")

    # ==========================================
    # 2. INSTITUTE CERTIFICATE (Appendix 2 Format)
    # ==========================================
    p_cert_hdr = doc.add_paragraph()
    p_cert_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cert_hdr.paragraph_format.space_before = Pt(12)
    p_cert_hdr.paragraph_format.space_after = Pt(12)
    r = p_cert_hdr.add_run("Major Project-I (01CE0716)\nDepartment of Computer Engineering\nFaculty of Engineering & Technology\nMarwadi University\nA.Y. 2026-27\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p_cert_title = doc.add_paragraph()
    p_cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cert_title.paragraph_format.space_after = Pt(20)
    r_ct = p_cert_title.add_run("CERTIFICATE")
    r_ct.font.name = 'Times New Roman'
    r_ct.font.size = Pt(18)
    r_ct.bold = True

    add_body_p(doc, "This is to certify that the project report submitted along with the project entitled TrustGuard: AI-Powered Flipkart Fake Review Detection System using Machine Learning & Deep Learning has been carried out by Jaykishan Kalariya (92301703102) and Maan Kalariya (92301703111) under my guidance in partial fulfilment for the degree of Bachelor of Technology in Computer Engineering, 7th Semester of Marwadi University, Rajkot during the academic year 2026-27.", space_after=40)

    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    sig_cells = sig_table.rows[0].cells
    sig_cells[0].text = "\n\n______________________"
    sig_cells[1].text = "\n\n______________________"
    
    sig_cells2 = sig_table.rows[1].cells
    sig_cells2[0].text = "Prof. Charmi Vora"
    sig_cells2[1].text = "Prof. (Dr.) Krunal Vaghela"
    
    sig_cells3 = sig_table.rows[2].cells
    sig_cells3[0].text = "Assistant Professor\nDepartment of Computer Engineering"
    sig_cells3[1].text = "Associate Dean & Head\nDepartment of Computer Engineering"

    for r_i in sig_table.rows:
        for c_i in r_i.cells:
            p = c_i.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    doc.add_page_break()

    # ==========================================
    # 3. DECLARATION (Appendix 3 Format)
    # ==========================================
    p_dec_hdr = doc.add_paragraph()
    p_dec_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dec_hdr.paragraph_format.space_before = Pt(12)
    p_dec_hdr.paragraph_format.space_after = Pt(12)
    r = p_dec_hdr.add_run("Major Project-I (01CE0716)\nDepartment of Computer Engineering\nFaculty of Engineering & Technology\nMarwadi University\nA.Y. 2026-27\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p_dec_t = doc.add_paragraph()
    p_dec_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dec_t.paragraph_format.space_after = Pt(20)
    r = p_dec_t.add_run("DECLARATION")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.bold = True

    add_body_p(doc, "We hereby declare that the Major Project-I (01CE0716) report submitted along with the Project entitled TrustGuard: AI-Powered Flipkart Fake Review Detection System using Machine Learning & Deep Learning submitted in partial fulfilment for the degree of Bachelor of Technology in Computer Engineering to Marwadi University, Rajkot, is a bonafide record of original project work carried out by us at Marwadi University under the supervision of Prof. Charmi Vora and that no part of this report has been directly copied from any students' reports or taken from any other source, without providing due reference.", space_after=30)

    dec_table = doc.add_table(rows=3, cols=2)
    dec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dec_cells = dec_table.rows[0].cells
    dec_cells[0].text = "Name of the Student"
    dec_cells[1].text = "Sign of Student"
    
    dec_cells1 = dec_table.rows[1].cells
    dec_cells1[0].text = "1. Jaykishan Kalariya (92301703102)"
    dec_cells1[1].text = "______________________"
    
    dec_cells2 = dec_table.rows[2].cells
    dec_cells2[0].text = "2. Maan Kalariya (92301703111)"
    dec_cells2[1].text = "______________________"

    for r_i in dec_table.rows:
        for c_i in r_i.cells:
            p = c_i.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    doc.add_page_break()

    # ==========================================
    # 4. ACKNOWLEDGEMENT
    # ==========================================
    add_heading_styled(doc, "Acknowledgement", level=1)
    add_body_p(doc, "We are deeply grateful and wish to express our sincere appreciation to all individuals and institutional mentors who have supported and guided us throughout the inception, design, development, and execution of this Major Project.")
    add_body_p(doc, "First and foremost, we express our profound gratitude to our internal project guide, Prof. Charmi Vora, Assistant Professor, Department of Computer Engineering. Her regular feedback, insightful technical suggestions, rigorous code evaluation, and encouragement at every development phase played a pivotal role in refining our natural language processing methodologies and model architectures.")
    add_body_p(doc, "We extend our deep gratitude to Prof. (Dr.) Krunal Vaghela, Associate Dean & Head of the Computer Engineering Department, Marwadi University, for granting access to high-performance computational infrastructure, software laboratories, and providing an encouraging academic environment essential for conducting extensive deep learning experiments.")
    add_body_p(doc, "We also express our appreciation to the entire faculty and technical staff of the Computer Engineering Department for imparting fundamental knowledge in Data Science, Machine Learning, Web Technologies, and Software Engineering throughout our academic journey.")
    add_body_p(doc, "Special thanks to our parents, family members, and peers for their unyielding moral support, patience, and motivation during long hours of research, model training, and project documentation.")
    add_body_p(doc, "Finally, we acknowledge the contributions of open-source research communities, Kaggle dataset creators, and the maintainers of PyTorch, Scikit-Learn, Flask, and Parse.bot API whose software packages made the realization of TrustGuard possible.")

    doc.add_page_break()

    # ==========================================
    # 5. ABSTRACT
    # ==========================================
    add_heading_styled(doc, "Abstract", level=1)
    add_body_p(doc, "In the modern digital economy, consumer purchasing behavior on e-commerce platforms such as Flipkart and Amazon is heavily dictated by customer product reviews and star rating aggregations. However, the rapidly rising phenomenon of Opinion Spam—manifested through computer-generated (CG) fake positive reviews, bot-driven rating inflation, and paid promotional spam—severely compromises consumer trust, deceives buyers into purchasing low-quality items, and inflicts financial losses. Manual inspection of thousands of weekly customer reviews is humanly impossible and highly error-prone, creating an urgent demand for automated, real-time Machine Learning and Natural Language Processing (NLP) verification frameworks.")
    add_body_p(doc, "This project presents TrustGuard, an end-to-end AI-powered fake review detection and product authenticity evaluation system designed specifically for Flipkart. The system integrates advanced NLP feature engineering, classical Machine Learning classifiers, and a custom PyTorch Bi-directional Long Short-Term Memory (BiLSTM) Deep Learning neural network.")
    add_body_p(doc, "The model is trained on a benchmark Kaggle dataset comprising 40,432 labeled review records (20,216 Computer-Generated fake reviews and 20,216 Original real customer reviews). Feature extraction combines Term Frequency-Inverse Document Frequency (TF-IDF) Word Unigrams & Bigrams (max 15,000 features), Character 2-4 N-Grams (max 10,000 features), and normalized behavioral stylistic metadata (ALL-CAPS word ratio, exclamation mark density, unique vocabulary diversity ratio, and review length). Multi-model benchmarking evaluates Multinomial Naive Bayes (90.96% accuracy), Logistic Regression (94.99% accuracy), Linear Support Vector Machine (95.71% accuracy), and the PyTorch BiLSTM Deep Learning network, which achieves the highest classification accuracy of 96.40% with an F1-Score of 96.38%.")
    add_body_p(doc, "To bridge theoretical machine learning with real-world e-commerce usage, TrustGuard incorporates the Parse.bot Flipkart API alongside an in-memory caching mechanism to extract live customer reviews from Flipkart product URLs. The system is deployed as a responsive Glassmorphic Flask web application that calculates real-time Product Trust Index percentage scores, renders Chart.js visual analytics, highlights suspicious promotional spam flags, and responds in under 0.02 seconds across laptops, tablets, and smartphones.")

    doc.add_page_break()

    # ==========================================
    # 6. LIST OF FIGURES & TABLES & ABBREVIATIONS
    # ==========================================
    add_heading_styled(doc, "List of Figures", level=1)
    fig_headers = ["Figure No.", "Figure Description", "Page No."]
    fig_data = [
        ["Fig 1.1", "Agile Software Development Lifecycle Model", "6"],
        ["Fig 2.1", "Data Pipeline & Scraper Process Flowchart", "11"],
        ["Fig 3.1", "TrustGuard Overall System Architecture Diagram", "16"],
        ["Fig 3.2", "Data Flow Diagram (Level 0 Context Diagram)", "17"],
        ["Fig 3.3", "Data Flow Diagram (Level 1 Process Breakdown)", "18"],
        ["Fig 3.4", "PyTorch BiLSTM Deep Learning Network Architecture", "20"],
        ["Fig 4.1", "TrustGuard Home Dashboard & Bulk Review Inspector", "30"],
        ["Fig 4.2", "Live Flipkart URL Inspector & Trust Score Ring", "31"],
        ["Fig 4.3", "Chart.js Authenticity Doughnut & Rating Breakdown Graphs", "32"],
        ["Fig 4.4", "Model Performance Benchmark Modal", "33"]
    ]
    create_table_styled(doc, fig_headers, fig_data, [1.2, 4.5, 0.8])

    doc.add_page_break()

    add_heading_styled(doc, "List of Tables", level=1)
    tbl_headers = ["Table No.", "Table Title", "Page No."]
    tbl_data = [
        ["Table 1.1", "Summary of Literature Review and Findings", "4"],
        ["Table 2.1", "Selection of Tools, Libraries & Technologies", "14"],
        ["Table 2.2", "Hardware & Software System Specifications", "15"],
        ["Table 3.1", "Database Schema & Serialized Model Data Structures", "19"],
        ["Table 4.1", "Model Performance Benchmark & Accuracy Comparison", "28"],
        ["Table 4.2", "Comprehensive Test Cases and Execution Results (TC-01 to TC-18)", "29"],
        ["Table 5.1", "Problems Encountered and Implemented Solutions", "37"]
    ]
    create_table_styled(doc, tbl_headers, tbl_data, [1.2, 4.5, 0.8])

    doc.add_page_break()

    add_heading_styled(doc, "Abbreviations", level=1)
    abb_headers = ["Abbreviation", "Full Expansion / Description"]
    abb_data = [
        ["AI", "Artificial Intelligence"],
        ["NLP", "Natural Language Processing"],
        ["ML", "Machine Learning"],
        ["DL", "Deep Learning"],
        ["BiLSTM", "Bi-directional Long Short-Term Memory"],
        ["SVM", "Support Vector Machine"],
        ["TF-IDF", "Term Frequency - Inverse Document Frequency"],
        ["CG", "Computer-Generated Fake Reviews"],
        ["OR", "Original Real Customer Reviews"],
        ["API", "Application Programming Interface"],
        ["REST", "Representational State Transfer"],
        ["MVC", "Model-View-Controller Architectural Pattern"],
        ["DFD", "Data Flow Diagram"],
        ["ERD", "Entity-Relationship Diagram"],
        ["RBAC", "Role-Based Access Control"],
        ["URL", "Uniform Resource Locator"],
        ["HTML", "HyperText Markup Language"],
        ["CSS", "Cascading Style Sheets"],
        ["JS", "JavaScript (ES6+)"],
        ["SDLC", "Software Development Life Cycle"]
    ]
    create_table_styled(doc, abb_headers, abb_data, [2.0, 4.5])

    doc.add_page_break()

    # ==========================================
    # CHAPTER 1.0 INTRODUCTION TO PROJECT & MANAGEMENT
    # ==========================================
    add_heading_styled(doc, "1.0 INTRODUCTION TO PROJECT AND PROJECT MANAGEMENT", level=1)
    
    add_heading_styled(doc, "1.1 Project Summary", level=2)
    add_body_p(doc, "In the rapidly expanding global e-commerce industry, customer product reviews represent the single most powerful factor driving online sales and buyer decision-making. Popular Indian e-commerce portals such as Flipkart host millions of products spanning consumer electronics, fashion, home appliances, and personal care. Prospective buyers rely on customer feedback to assess product quality, durability, delivery speed, and seller authenticity prior to financial transactions.")
    add_body_p(doc, "However, this reliance has incentivized unethical sellers, third-party marketing agencies, and automated bot networks to inject fake reviews into e-commerce product pages. These fake reviews manifest in two primary forms: (1) Computer-Generated (CG) positive fake reviews designed to artificially inflate product ratings, and (2) Paid promotional spam containing repetitive, generic praise without substantive product information. This systemic manipulation results in severe consumer financial loss, degraded trust in e-commerce platforms, and unfair market disadvantage for genuine sellers who deliver quality products without resorting to deceptive tactics.")
    add_body_p(doc, "To address this critical challenge, this project introduces TrustGuard—an intelligent, automated fake review detection and product authenticity auditing framework built specifically for Flipkart product pages. TrustGuard bridges natural language processing (NLP), statistical text modeling, classical machine learning, and deep neural networks into a unified web application.")
    add_body_p(doc, "The underlying machine learning models are trained on a comprehensive Kaggle dataset containing 40,432 labeled review instances (20,216 CG fake reviews and 20,216 OR real customer reviews). The feature extraction pipeline extracts a hybrid matrix combining TF-IDF Word N-Grams (unigrams and bigrams), TF-IDF Character 2-4 N-Grams, and normalized stylistic behavioral metadata (all-caps word ratio, exclamation mark count, unique vocabulary diversity ratio, review length, and star rating).")
    add_body_p(doc, "Four distinct model architectures were implemented and benchmarked: Multinomial Naive Bayes (90.96% accuracy), Logistic Regression (94.99% accuracy), Linear Support Vector Machine (95.71% accuracy), and a PyTorch Bi-directional Long Short-Term Memory (BiLSTM) Deep Learning neural network, which achieved the highest classification accuracy of 96.40% and an F1-Score of 96.38%. The models are serialized into binary weights (.pkl and .pt) for instant in-memory loading (< 0.01s). Integrated with the Parse.bot Flipkart API, an in-memory URL caching layer, and a responsive Glassmorphic Flask web interface, TrustGuard enables users to inspect live Flipkart product links or bulk-paste review text, rendering instant Trust Index percentage scores and Chart.js visual analytics across desktop, tablet, and mobile devices.")

    add_heading_styled(doc, "1.2 Purpose", level=2)
    add_body_p(doc, "The primary purpose of developing TrustGuard is to establish an automated, objective, and mathematically sound protection mechanism against online review manipulation. The detailed purposes of this project include:")
    add_bullet_p(doc, "Traditional human moderation is slow, subjective, and unable to scale to millions of daily customer reviews. TrustGuard automates review extraction and instant authenticity classification.", bold_prefix="1. Automated E-Commerce Fraud Prevention: ")
    add_bullet_p(doc, "Classical sentiment analyzers only evaluate positive or negative polarity. TrustGuard combines textual n-grams with stylistic metadata (all-caps shouting, exclamation density, repetitive vocabulary) to detect structural spam patterns.", bold_prefix="2. Hybrid NLP Feature Extraction: ")
    add_bullet_p(doc, "Consumers often struggle to assess product credibility from raw average star ratings. TrustGuard aggregates individual review classifications into an intuitive single-number Trust Index Percentage = (Real Reviews / Total Reviews) * 100.", bold_prefix="3. Consumer Decision Support Metric: ")
    add_bullet_p(doc, "Direct web scraping of Flipkart pages is frequently blocked by security firewalls. TrustGuard integrates the Parse.bot Flipkart API backed by an in-memory URL caching layer to ensure 100% reliable live data fetching.", bold_prefix="4. Resilient Live Data Scraping: ")
    add_bullet_p(doc, "To deliver a sub-second response time (< 0.02 seconds) and responsive touch-friendly interface accessible across smartphones, tablets, laptops, and desktop computers.", bold_prefix="5. Cross-Device High-Performance Access: ")

    add_heading_styled(doc, "1.3 Objective", level=2)
    add_body_p(doc, "The project is guided by specific quantitative and qualitative engineering objectives:")
    add_bullet_p(doc, "To preprocess, clean, and stratify a 40,432-row e-commerce review dataset into 80% training (32,346 samples) and 20% testing (8,086 samples) subsets.", bold_prefix="Objective 1 (Dataset Curation & Preprocessing): ")
    add_bullet_p(doc, "To construct a hybrid NLP vectorization pipeline combining 15,000 TF-IDF Word N-Grams, 10,000 TF-IDF Char N-Grams, and 7 scaled stylistic metadata attributes using Scikit-Learn StandardScaler.", bold_prefix="Objective 2 (NLP Feature Engineering): ")
    add_bullet_p(doc, "To train and benchmark Multinomial Naive Bayes, Logistic Regression, Linear SVM, and a PyTorch Bi-directional LSTM neural network, targeting a benchmark classification accuracy exceeding 95%.", bold_prefix="Objective 3 (Model Benchmarking): ")
    add_bullet_p(doc, "To implement live review scraping from Flipkart product URLs via Parse.bot API with an in-memory caching dictionary to conserve monthly API credits.", bold_prefix="Objective 4 (Live Scraper & Cache System): ")
    add_bullet_p(doc, "To build a responsive Flask REST API server (/api/analyze-url, /api/analyze-bulk, /api/models, /ping) supporting cross-device network host binding (0.0.0.0).", bold_prefix="Objective 5 (Flask Web Backend): ")
    add_bullet_p(doc, "To design an intuitive Glassmorphism single-page frontend interface featuring Chart.js doughnut charts, rating breakdown bar charts, review feed filters, and model switching dropdowns.", bold_prefix="Objective 6 (Interactive Visual Dashboard): ")

    add_heading_styled(doc, "1.4 Scope", level=2)
    add_body_p(doc, "System Capabilities (In Scope):", bold_prefix="1.4.1 In-Scope System Capabilities: ")
    add_bullet_p(doc, "Extracts up to 30 live customer reviews directly from any valid Flipkart product URL across multiple pages.")
    add_bullet_p(doc, "Allows users to bulk copy-paste raw review text directly from Flipkart product pages for instant analysis.")
    add_bullet_p(doc, "Evaluates individual review text and generates a REAL vs FAKE prediction with confidence percentage.")
    add_bullet_p(doc, "Calculates an overall Product Trust Index Percentage = (Real Reviews / Total Reviews) * 100.")
    add_bullet_p(doc, "Detects and flags specific spam triggers: Excessive Capitalization, Repeated Exclamations, Repetitive Vocabulary, and Generic Promotional Phrases.")
    add_bullet_p(doc, "Renders interactive Chart.js doughnut charts for review authenticity split and bar charts for rating distribution.")
    add_bullet_p(doc, "Provides real-time model switching in the web header between PyTorch BiLSTM, Linear SVM, Logistic Regression, and Naive Bayes.")

    add_body_p(doc, "System Limitations (Out of Scope):", bold_prefix="1.4.2 Out-of-Scope System Limitations: ")
    add_bullet_p(doc, "Does not analyze regional non-English languages (Hindi, Hinglish, Tamil) in the current release.")
    add_bullet_p(doc, "Does not include native iOS or Android mobile application binaries (.apk / .ipa); operates via mobile web browsers.")
    add_bullet_p(doc, "Does not inspect product images or video attachments in customer reviews.")

    add_body_p(doc, "Future Scope:", bold_prefix="1.4.3 Future Scope & Enhancements: ")
    add_bullet_p(doc, "Integration with transformer architectures (BERT, RoBERTa, DeBERTa) for deep semantic context.")
    add_bullet_p(doc, "Development of a Chrome Browser Extension to overlay Trust Badges directly on Flipkart shopping pages.")
    add_bullet_p(doc, "Expansion of scraping adapters to support Amazon India, Myntra, Meesho, and Nykaa.")

    add_heading_styled(doc, "1.5 Technology and Literature Review", level=2)
    add_body_p(doc, "Literature Review Summary Table:", bold_prefix="1.5.1 Summary of Literature Review and Findings: ")
    add_body_p(doc, "Table 1.1 provides a structured synthesis of key research papers in opinion spam detection, time series sentiment modeling, and neural text classification that informed the development of TrustGuard.")

    lit_headers = ["Ref No.", "Author & Year", "Model / Architecture", "Dataset Used", "Accuracy (%)", "Key Findings & Insights"]
    lit_data = [
        ["1", "Jindal & Liu (2008)", "Logistic Regression & Rules", "Amazon Reviews", "84.0%", "Identified duplicate reviews as primary spam signal."],
        ["2", "Ott et al. (2011)", "Linear SVM & Naive Bayes", "Hotel Reviews (Deceptive)", "89.8%", "Proved humans achieve only 57% accuracy vs 89.8% ML."],
        ["3", "Crawford et al. (2015)", "Ensemble ML Models", "E-Commerce Reviews", "88.5%", "Demonstrated metadata + text n-grams beats text-only."],
        ["4", "Zhang et al. (2020)", "Bi-directional LSTM (BiLSTM)", "Yelp & Amazon Data", "94.2%", "Proved BiLSTM captures past/future context superior to RNN."],
        ["5", "TrustGuard (Our Work)", "PyTorch BiLSTM + Hybrid NLP", "Kaggle 40k Dataset", "96.40%", "Combines BiLSTM, TF-IDF n-grams, and scaled metadata."]
    ]
    create_table_styled(doc, lit_headers, lit_data, [0.8, 1.5, 1.8, 1.3, 1.0, 2.5])

    add_heading_styled(doc, "1.6 Project Planning and Scheduling", level=2)
    add_heading_styled(doc, "1.6.1 Development Approach Flowchart", level=3)
    add_body_p(doc, "Fig 1.1 illustrates the iterative Agile software development lifecycle used throughout the 28-day project timeline.")

    # Flowchart Visual Box
    add_formula_box(doc, "Fig 1.1: Agile Software Development Lifecycle Flowchart", 
        "+--------------------------------------------------------------+\n"
        "|  1. Requirement Analysis (Dataset Curation & Scraper Setup)  |\n"
        "+------------------------------+------------------------------+\n"
        "                               |\n"
        "                               v\n"
        "+--------------------------------------------------------------+\n"
        "|  2. NLP Feature Engineering (TF-IDF Word/Char + Metadata)   |\n"
        "+------------------------------+------------------------------+\n"
        "                               |\n"
        "                               v\n"
        "+--------------------------------------------------------------+\n"
        "|  3. Model Training & Tuning (PyTorch BiLSTM, SVM, LogReg)   |\n"
        "+------------------------------+------------------------------+\n"
        "                               |\n"
        "                               v\n"
        "+--------------------------------------------------------------+\n"
        "|  4. Flask REST API Backend & Parse.bot Live Scraper Cache    |\n"
        "+------------------------------+------------------------------+\n"
        "                               |\n"
        "                               v\n"
        "+--------------------------------------------------------------+\n"
        "|  5. Responsive Glassmorphic UI & Chart.js Visual Analytics   |\n"
        "+------------------------------+------------------------------+\n"
        "                               |\n"
        "                               v\n"
        "+--------------------------------------------------------------+\n"
        "|  6. Testing (18 Test Cases), Mobile Validation & Cloud Deploy|\n"
        "+--------------------------------------------------------------+"
    )

    doc.add_page_break()

    # ==========================================
    # CHAPTER 2.0 SYSTEM ANALYSIS
    # ==========================================
    add_heading_styled(doc, "2.0 SYSTEM ANALYSIS", level=1)
    
    add_heading_styled(doc, "2.1 Study of Current System", level=2)
    add_body_p(doc, "In existing e-commerce systems, customer product review verification is conducted using two primary approaches: manual human moderation and rule-based keyword blacklists.")
    add_body_p(doc, "Manual Human Moderation: In this traditional approach, human content moderators manually read reported reviews to spot spam or abusive language. However, human moderators cannot scale to process tens of thousands of reviews posted daily. Furthermore, academic research (Ott et al., 2011) has proven that humans achieve only 57% accuracy when evaluating deceptive positive reviews, as computer-generated fake reviews are specifically crafted to sound natural and persuasive.")
    add_body_p(doc, "Rule-Based Keyword Blacklists: Many automated web scripts look exclusively for blacklisted profanity, offensive keywords, or external website URLs. While effective at blocking overt vulgarity, keyword filters fail completely against sophisticated fake reviews (e.g. 'This product is amazing! Fast delivery, highly recommended!'). Because these fake reviews contain grammatically valid, highly positive words, keyword blacklists flag them as 100% genuine.")

    add_heading_styled(doc, "2.2 Activity Diagram of Data Pipeline", level=2)
    add_body_p(doc, "Fig 2.1 displays the ten-stage linear activity and data pipeline process flow implemented in TrustGuard.")

    add_formula_box(doc, "Fig 2.1: Activity Diagram of the Data Pipeline Process",
        "[START]\n"
        "   |\n"
        "   v\n"
        "1. Input Acquisition (Parse.bot Live Scraping / Bulk Paste Text)\n"
        "   |\n"
        "   v\n"
        "2. Input Validation (Domain Matching: flipkart.com / Path: /p/)\n"
        "   |\n"
        "   v\n"
        "3. Text Preprocessing (Lowercasing, URL/HTML Stripping, Noise Removal)\n"
        "   |\n"
        "   v\n"
        "4. Feature Engineering (TF-IDF Word 1-2 N-Grams + Char 2-4 N-Grams)\n"
        "   |\n"
        "   v\n"
        "5. Metadata Extraction (Caps Ratio, Exclamation Density, Unique Diversity)\n"
        "   |\n"
        "   v\n"
        "6. StandardScaler Scaling (Standardized Metadata Alignment)\n"
        "   |\n"
        "   v\n"
        "7. PyTorch BiLSTM / ML Model Matrix Inference (< 0.02s)\n"
        "   |\n"
        "   v\n"
        "8. Suspicion Flag Evaluator (Promotional Phrases, Caps, Exclamations)\n"
        "   |\n"
        "   v\n"
        "9. Trust Index % Calculation = (Real Reviews / Total Reviews) * 100\n"
        "   |\n"
        "   v\n"
        "10. Visual Analytics Output (Glassmorphism Dashboard & Chart.js Graphs)\n"
        "   |\n"
        "   v\n"
        "[END]"
    )

    add_heading_styled(doc, "2.8 Selection of Tools and Technologies", level=2)
    sw_headers = ["Category", "Tool / Library", "Purpose & Technical Specification"]
    sw_data = [
        ["Language", "Python 3.11", "Primary programming runtime for ML models, scraper, and Flask web server."],
        ["Deep Learning", "PyTorch (torch 2.13)", "5-layer Bi-directional LSTM neural network for sequence text classification."],
        ["Machine Learning", "Scikit-Learn 1.4", "Linear SVM, Logistic Regression, Naive Bayes, TF-IDF Vectorizers, StandardScaler."],
        ["Web Framework", "Flask 3.0", "Python micro web framework serving REST API endpoints (/api/analyze-url)."],
        ["Scraping API", "Parse.bot API", "Live scraping API extracting up to 30 customer reviews per Flipkart URL."],
        ["Visualization", "Chart.js 4.4 & SVG", "Renders Review Authenticity Doughnut Chart and Rating Breakdown Bar Chart."]
    ]
    create_table_styled(doc, sw_headers, sw_data, [1.5, 1.8, 3.2])

    add_heading_styled(doc, "2.9 Mathematical & Algorithmic Formulations", level=2)
    add_body_p(doc, "1. Term Frequency - Inverse Document Frequency (TF-IDF):", bold_prefix="2.9.1 TF-IDF Formulation: ")
    add_body_p(doc, "TF-IDF evaluates the relative importance of unigrams, bigrams, and character n-grams across the review corpus. Formula:")
    
    add_formula_box(doc, "Formula 1: Term Frequency - Inverse Document Frequency",
        "TF-IDF(t, d, D) = TF(t, d) * [ log( (1 + |D|) / (1 + |{d in D : t in d}|) ) + 1 ]"
    )

    add_body_p(doc, "2. StandardScaler Z-Score Normalization:", bold_prefix="2.9.2 Z-Score Metadata Normalization: ")
    add_body_p(doc, "Stylistic metadata attributes (Caps Ratio, Exclamation Count, Review Length) are normalized to zero mean and unit variance:")
    
    add_formula_box(doc, "Formula 2: Z-Score Standardization",
        "z = (x - mu) / sigma"
    )

    add_body_p(doc, "3. Linear Support Vector Machine (Linear SVM):", bold_prefix="2.9.3 Linear SVM Hyperplane & Decision Function: ")
    add_body_p(doc, "Linear SVM seeks the maximum-margin hyperplane w^T * X + b = 0 separating REAL and FAKE feature vectors:")
    
    add_formula_box(doc, "Formula 3: Linear SVM Objective & Sigmoidal Distance Transformation",
        "min_{w,b} (1/2) * ||w||^2 + C * SUM( max(0, 1 - y_i * (w^T * x_i + b)) )\n\n"
        "P(Fake) = 1 / ( 1 + exp( -1.2 * (w^T * X + b) ) )"
    )

    add_body_p(doc, "4. PyTorch Bi-directional LSTM (BiLSTM) Cell State Equations:", bold_prefix="2.9.4 BiLSTM Neural Network Architecture Equations: ")
    add_body_p(doc, "The BiLSTM model processes sequences in both forward (left-to-right) and backward (right-to-left) directions. At time step t, the cell state updates via Input Gate (i_t), Forget Gate (f_t), and Output Gate (o_t):")
    
    add_formula_box(doc, "Formula 4: PyTorch Bi-directional LSTM Neural Equations",
        "f_t = sigmoid( W_f * [h_{t-1}, x_t] + b_f )    [Forget Gate]\n"
        "i_t = sigmoid( W_i * [h_{t-1}, x_t] + b_i )    [Input Gate]\n"
        "C_tilde_t = tanh( W_c * [h_{t-1}, x_t] + b_c ) [Candidate Cell State]\n"
        "C_t = f_t (*) C_{t-1} + i_t (*) C_tilde_t       [Cell State Update]\n"
        "o_t = sigmoid( W_o * [h_{t-1}, x_t] + b_o )    [Output Gate]\n"
        "h_t = o_t (*) tanh( C_t )                       [Hidden Output State]\n\n"
        "h_BiLSTM = [ h_forward_t ; h_backward_t ]      [Bidirectional Concatenation]\n\n"
        "P(Fake) = sigmoid( W_2 * ReLU( W_1 * MeanPool( h_BiLSTM ) ) )"
    )

    add_body_p(doc, "5. Product Trust Index Percentage Formula:", bold_prefix="2.9.5 Product Trust Index Calculation: ")
    
    add_formula_box(doc, "Formula 5: Product Trust Index Score",
        "Trust Index % = ( Count(REAL Reviews) / Total Reviews Analyzed ) * 100%"
    )

    doc.add_page_break()

    # ==========================================
    # CHAPTER 3.0 SYSTEM DESIGN
    # ==========================================
    add_heading_styled(doc, "3.0 SYSTEM DESIGN", level=1)
    
    add_heading_styled(doc, "3.1 System Architecture Diagram", level=2)
    add_body_p(doc, "Fig 3.1 illustrates the overall 3-tier system architecture and data pipeline of TrustGuard.")

    add_formula_box(doc, "Fig 3.1: Overall System Architecture of TrustGuard",
        "+-------------------------------------------------------------------------+\n"
        "|                 DATA SOURCE (Flipkart URL / Bulk Text)                  |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  APPLICATION TIER: Parse.bot Scraper & In-Memory URL_CACHE Hash Table   |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  NLP FEATURE ENGINE: TF-IDF Word (1-2) + Char (2-4) + Scaled Metadata   |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  CLASSIFICATION ENGINE: PyTorch BiLSTM | Linear SVM | LogReg | Naive Bayes|\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  PRESENTATION TIER: Trust Ring SVG | Chart.js Analytics | Review Feed   |\n"
        "+-------------------------------------------------------------------------+"
    )

    add_heading_styled(doc, "3.2 Data Flow Diagrams (DFD Level 0 & Level 1)", level=2)
    add_body_p(doc, "Fig 3.2: DFD Level 0 Context Diagram", bold_prefix="DFD Level 0 Context Diagram: ")
    
    add_formula_box(doc, "Fig 3.2: Data Flow Diagram (Level 0 Context Diagram)",
        " [ Shopper / User ] ---> ( Flipkart URL / Text ) ---> [ 0. TRUSTGUARD ENGINE ]\n"
        "                                                            |\n"
        " [ Shopper / User ] <--- ( Trust Score & Analytics ) <------+"
    )

    add_body_p(doc, "Fig 3.3: DFD Level 1 Detailed Process Breakdown", bold_prefix="DFD Level 1 Detailed Process Breakdown: ")

    add_formula_box(doc, "Fig 3.3: Data Flow Diagram (Level 1 Process Breakdown)",
        "[ User ] -> (1.0 Scraping & Cache) -> (2.0 Preprocessing) -> (3.0 Feature Matrix)\n"
        "                                                                     |\n"
        "[ Dashboard ] <- (5.0 Output Generator) <- (4.0 BiLSTM/ML Classifier) <+"
    )

    add_heading_styled(doc, "3.4 Deep Learning Architecture (PyTorch BiLSTM)", level=2)
    add_body_p(doc, "Fig 3.4 details the internal neural network layer configuration of the PyTorch Bi-directional LSTM model.")

    add_formula_box(doc, "Fig 3.4: PyTorch BiLSTM Network Architecture",
        "+-------------------------------------------------------------------------+\n"
        "|  Input Layer (Word Token Integer Sequence Vector, max_len=120)          |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Embedding Layer (vocab_size=20,000, embed_dim=128, padding_idx=0)       |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Bidirectional LSTM Layer (hidden_dim=64, bidirectional=True, concat=128)|\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Mean Pooling Layer (Temporal Sequence Reduction over Time Dim)         |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Dense FC1 (128 -> 32) + ReLU Activation + Dropout(0.3)                 |\n"
        "+------------------------------------+------------------------------------\n"
        "                                     |\n"
        "                                     v\n"
        "+-------------------------------------------------------------------------+\n"
        "|  Dense FC2 Output (32 -> 1) + Sigmoid Activation ---> Probability P(Fake)|\n"
        "+-------------------------------------------------------------------------+"
    )

    add_heading_styled(doc, "3.5 Database & Data Structure Design", level=2)
    ds_headers = ["Data Structure", "Storage Format", "Description & Field Schema"]
    ds_data = [
        ["Model Weight Binaries", "Joblib / PyTorch (.pkl / .pt)", "Stores trained weights for Linear SVM, LogReg, Naive Bayes & BiLSTM."],
        ["Vectorizer Vocabularies", "Pickle (.pkl)", "Stores 15,000 word TF-IDF & 10,000 char TF-IDF vocabulary mappings."],
        ["StandardScaler Object", "Pickle (.pkl)", "Stores feature means and standard deviations for metadata scaling."],
        ["URL_CACHE Hash Table", "Python In-Memory Dict", "Stores key-value pairs {url: scraped_review_json} to save API credits."],
        ["Analyzed Review Payload", "JSON Object", "{product_name, url, model_used, summary, reviews: [{author, rating, prediction, confidence, flags}]}"]
    ]
    create_table_styled(doc, ds_headers, ds_data, [1.8, 2.2, 3.0])

    doc.add_page_break()

    # ==========================================
    # CHAPTER 4.0 IMPLEMENTATION & TESTING
    # ==========================================
    add_heading_styled(doc, "4.0 IMPLEMENTATION AND TESTING", level=1)
    
    add_heading_styled(doc, "4.1 Implementation Environment", level=2)
    env_headers = ["Component", "Software / Tool Used", "Version / Specifications"]
    env_data = [
        ["Operating System", "macOS Sonoma / Linux Ubuntu", "64-bit Architecture"],
        ["Language & Environment", "Python", "v3.11.8"],
        ["Deep Learning Framework", "PyTorch (torch)", "v2.13.0 (CPU/ARM64)"],
        ["Machine Learning Library", "Scikit-Learn", "v1.4.0"],
        ["Web Framework", "Flask", "v3.0.2"],
        ["Web Scraping API", "Parse.bot API", "Live Flipkart Scraper"],
        ["Frontend Libraries", "Chart.js & FontAwesome", "Chart.js v4.4, FontAwesome 6"]
    ]
    create_table_styled(doc, env_headers, env_data, [1.8, 2.7, 2.5])

    add_heading_styled(doc, "4.2 Experimental Model Benchmark Results", level=2)
    bm_headers = ["Model Architecture", "Model Type", "Accuracy Score (%)", "F1 Score (%)", "Feature Set Used"]
    bm_data = [
        ["BiLSTM (Bi-directional LSTM)", "Deep Learning", "96.40%", "96.38%", "Word Embeddings + Dense Neural Layers"],
        ["Linear Support Vector Machine", "Machine Learning", "95.71%", "95.70%", "TF-IDF Word + Char N-Grams + Scaled Meta"],
        ["Logistic Regression", "Machine Learning", "94.99%", "94.97%", "TF-IDF Word + Char N-Grams + Scaled Meta"],
        ["Multinomial Naive Bayes", "Machine Learning", "90.96%", "90.96%", "TF-IDF Word Matrix"]
    ]
    create_table_styled(doc, bm_headers, bm_data, [2.2, 1.3, 1.2, 1.2, 1.6])

    add_heading_styled(doc, "4.3 Comprehensive Test Cases & Execution Results", level=2)
    tc_headers = ["Test ID", "Module", "Test Condition", "Expected Output", "Actual Output", "Status"]
    tc_data = [
        ["TC-01", "Text Input", "Submit empty text", "Validation error shown", "Error displayed", "PASS"],
        ["TC-02", "Bulk Input", "Paste 10 real reviews", "Analyzes 10 items", "10 items analyzed", "PASS"],
        ["TC-03", "URL Input", "Valid Flipkart URL", "Fetches 30 live reviews", "30 reviews fetched", "PASS"],
        ["TC-04", "URL Input", "Non-Flipkart URL", "Rejects with alert", "Alert displayed", "PASS"],
        ["TC-05", "URL Input", "Duplicate URL query", "Returns from cache", "Returned in 0.001s", "PASS"],
        ["TC-06", "Model Select", "Switch to BiLSTM", "Recalculates with BiLSTM", "BiLSTM outputs shown", "PASS"],
        ["TC-07", "Model Select", "Switch to Linear SVM", "Recalculates with SVM", "SVM outputs shown", "PASS"],
        ["TC-08", "Flags", "Text with BUY BUY NOW", "Flags Promotional Phrase", "Flagged correctly", "PASS"],
        ["TC-09", "Flags", "Text with ALL CAPS", "Flags Excessive Caps", "Flagged correctly", "PASS"],
        ["TC-10", "Trust Ring", "All Real Reviews", "Trust Score = 100%", "Ring shows 100%", "PASS"],
        ["TC-11", "Trust Ring", "All Fake Reviews", "Trust Score = 0%", "Ring shows 0%", "PASS"],
        ["TC-12", "Charts", "Render Doughnut Chart", "Displays Real/Fake split", "Chart rendered", "PASS"],
        ["TC-13", "Charts", "Render Bar Chart", "Displays Rating breakdown", "Chart rendered", "PASS"],
        ["TC-14", "Filters", "Click Fake Only filter", "Hides real reviews", "Only fake shown", "PASS"],
        ["TC-15", "Modal", "Click Benchmark button", "Opens metrics table", "Modal displayed", "PASS"],
        ["TC-16", "Health Check", "GET /ping endpoint", "Returns status: active", "Status 200 OK", "PASS"],
        ["TC-17", "Responsiveness", "View on Mobile device", "Layout stacks cleanly", "Mobile layout active", "PASS"],
        ["TC-18", "Inference", "BiLSTM Prediction", "Evaluates in < 0.02s", "Sub-second response", "PASS"]
    ]
    create_table_styled(doc, tc_headers, tc_data, [0.8, 1.0, 1.7, 1.5, 1.3, 0.7])

    add_heading_styled(doc, "4.4 System Visual Figures & Screen Descriptions", level=2)
    add_body_p(doc, "Fig 4.1: TrustGuard Home Dashboard & Bulk Review Inspector - Displays single-page Glassmorphic layout with multi-tab inputs and one-word action buttons.")
    add_body_p(doc, "Fig 4.2: Live Flipkart URL Inspector & Trust Index Score Ring - Displays SVG circular score ring showing product trust percentage.")
    add_body_p(doc, "Fig 4.3: Chart.js Authenticity Doughnut & Rating Breakdown Bar Graphs - Displays interactive charts rendering real/fake distribution and star rating breakdown.")
    add_body_p(doc, "Fig 4.4: Model Performance Benchmark Modal - Displays interactive comparison modal listing metrics for BiLSTM, Linear SVM, LogReg, and Naive Bayes.")

    doc.add_page_break()

    # ==========================================
    # CHAPTER 5.0 CONCLUSION & EXPANDED FUTURE SCOPE
    # ==========================================
    add_heading_styled(doc, "5.0 CONCLUSION AND FUTURE ENHANCEMENTS", level=1)
    
    add_heading_styled(doc, "5.1 Overall Analysis of Project Viabilities", level=2)
    add_body_p(doc, "The TrustGuard project has been successfully designed, implemented, and tested as an AI-powered fake review detection system. The practical combination of TF-IDF word/char n-grams, metadata scaling, and a PyTorch BiLSTM neural network achieves an outstanding 96.40% classification accuracy on 40,432 e-commerce reviews.")

    add_heading_styled(doc, "5.2 Problems Encountered and Possible Solutions", level=2)
    prob_sol_headers = ["Problem Encountered", "Technical Cause", "Solution Implemented"]
    prob_sol_data = [
        ["Flipkart Anti-Scraping", "Direct requests return 403 Forbidden", "Integrated Parse.bot Flipkart API."],
        ["API 200 Credit Limit", "Scraper credits limited per month", "Added in-memory URL caching layer."],
        ["Render App Sleep Mode", "Free web service spins down after 15m", "Added /ping route for UptimeRobot pings."],
        ["Mobile UI Clutter", "Long button texts wrapped clumsily on phones", "Redesigned UI with one-word action buttons."]
    ]
    create_table_styled(doc, prob_sol_headers, prob_sol_data, [1.8, 2.5, 2.7])

    add_heading_styled(doc, "5.5 Expanded Future Enhancements", level=2)
    add_body_p(doc, "1. Fine-Tuning Transformer Architectures (BERT / DeBERTa / RoBERTa):", bold_prefix="5.5.1 Fine-Tuning Transformer Architectures: ")
    add_body_p(doc, "While the PyTorch BiLSTM model achieves an impressive 96.40% accuracy, future iterations will explore fine-tuning pre-trained Transformer language models such as DeBERTa-v3 and RoBERTa-large. Transformers utilize multi-head self-attention mechanisms to capture subtle contextual nuances, sarcasm, and indirect spam phrases across long review texts.")

    add_body_p(doc, "2. Real-Time Chrome Browser Extension Development:", bold_prefix="5.5.2 Real-Time Chrome Extension Development: ")
    add_body_p(doc, "To provide direct utility to online shoppers during active browsing, a Manifest V3 Chrome Extension will be developed. The extension will automatically parse Flipkart product pages, send review payloads to the TrustGuard REST API, and inject a floating Trust Badge directly beside the product title on Flipkart.com.")

    add_body_p(doc, "3. Multi-Platform E-Commerce Scraper Expansion:", bold_prefix="5.5.3 Multi-Platform E-Commerce Expansion: ")
    add_body_p(doc, "While current data collection targets Flipkart, future versions will incorporate dedicated scraping adapters for Amazon India, Myntra, Meesho, and Nykaa, creating a universal e-commerce authenticity auditor.")

    add_body_p(doc, "4. Multilingual & Code-Switched Sentiment Analysis:", bold_prefix="5.5.4 Multilingual Sentiment Analysis (Hinglish & Regional Languages): ")
    add_body_p(doc, "Indian e-commerce reviews frequently feature code-switched text (Hinglish, e.g. 'Bahut achha product hai, super fast delivery!'). Future NLP pipelines will incorporate mBERT (Multilingual BERT) to classify regional language customer reviews.")

    add_body_p(doc, "5. Explainable AI (XAI) Integration via SHAP & LIME:", bold_prefix="5.5.5 Explainable AI (XAI) Integration: ")
    add_body_p(doc, "To enhance user transparency, future updates will incorporate SHAP (SHapley Additive exPlanations) and LIME (Local Interpretable Model-agnostic Explanations) to highlight the exact words and features that triggered a fake review prediction.")

    doc.add_page_break()

    # ==========================================
    # REFERENCES / BIBLIOGRAPHY
    # ==========================================
    add_heading_styled(doc, "References / Bibliography", level=1)
    add_heading_styled(doc, "Books:", level=2)
    add_body_p(doc, "1. Jurafsky, D., & Martin, J. H. (2023). Speech and Language Processing (3rd ed. draft). Pearson Education.")
    add_body_p(doc, "2. Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.")

    add_heading_styled(doc, "Online References:", level=2)
    add_body_p(doc, "1. PyTorch Documentation & Tutorials: https://pytorch.org/docs/")
    add_body_p(doc, "2. Scikit-Learn Machine Learning Documentation: https://scikit-learn.org/")
    add_body_p(doc, "3. Parse.bot Flipkart API Documentation: https://parse.bot/")
    add_body_p(doc, "4. Flask Micro Web Framework: https://flask.palletsprojects.com/")

    add_heading_styled(doc, "Research Papers:", level=2)
    add_body_p(doc, "1. Jindal, N., & Liu, B. (2008). Opinion spam and analysis. Proceedings of the 2008 International Conference on Web Search and Data Mining (WSDM '08), 219–230.")
    add_body_p(doc, "2. Crawford, M., Khoshgoftaar, T. M., Prusa, J. D., Richter, A. N., & Najada, H. (2015). Survey of review spam detection using machine learning techniques. Journal of Big Data, 2(1), 1–24.")

    doc.save("MARWADI_UNIVERSITY_PROJECT_REPORT.docx")
    print("Successfully generated MASTER MARWADI_UNIVERSITY_PROJECT_REPORT.docx with clean formulas, visual flowcharts, streamlined tables, and expanded future scope!")

if __name__ == '__main__':
    generate_master_report()
